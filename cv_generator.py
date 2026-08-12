from __future__ import annotations

from datetime import date, datetime
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from database import Database
from cv_data import build_cv_data, full_text, pretty_date

NAVY = RGBColor(31, 78, 121)
DARK = RGBColor(45, 45, 45)


def add_complete_text(doc: Document, text: Any, *, bullet: bool = False) -> None:
    value = full_text(text)
    lines = value.splitlines() or [value]
    for line in lines:
        if bullet:
            add_bullet(doc, line)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.add_run(line)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY
    p_pr = p._p.get_or_add_pPr()
    bottom = OxmlElement('w:pBdr')
    border = OxmlElement('w:bottom')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '8')
    border.set(qn('w:space'), '1')
    border.set(qn('w:color'), '1F4E79')
    bottom.append(border)
    p_pr.append(bottom)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.add_run(text)


def generate_cv(db: Database, output_path: str | Path, options: dict[str, Any] | None = None) -> Path:
    options = options or {}
    output_path = Path(output_path)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles['Normal'].font.name = 'Aptos'
    styles['Normal'].font.size = Pt(9.5)
    styles['Normal'].font.color.rgb = DARK

    data = build_cv_data(db, options)
    profile = data["profile"]
    sections = data["sections"]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(profile.get('full_name', ''))
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(profile.get('title', ''))
    r.bold = True
    r.font.size = Pt(11)

    contact = " | ".join(x for x in [profile.get('agency'), profile.get('email'), profile.get('phone')] if x)
    p = doc.add_paragraph(contact)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)

    if options.get('summary', True) and profile.get('summary'):
        add_section_heading(doc, "Professional Summary")
        add_complete_text(doc, profile['summary'])

    if options.get('core_training', True):
        core = sections.get('core_training', [])
        if core:
            add_section_heading(doc, "Core Training")
            for x in core:
                attended = pretty_date(full_text(x.get('attended_date')))
                provider = f" ({full_text(x.get('provider'))})" if x.get('provider') else ''
                add_bullet(doc, f"{full_text(x.get('course_name'))}{provider}, {attended}".strip(', '))

    if options.get('employment', True):
        rows = sections.get('employment', [])
        if rows:
            add_section_heading(doc, "Work Experience")
            for x in rows:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(1)
                period = " to ".join(filter(None, [pretty_date(x.get('start_date'), True), pretty_date(x.get('end_date'), True)]))
                r = p.add_run(f"{period} - {x.get('employer')} ({x.get('title')})")
                r.bold = True
                add_complete_text(doc, x.get('description'), bullet=True)

    if options.get('teaching', True):
        rows = sections.get('teaching', [])
        if rows:
            add_section_heading(doc, "Teaching Experience")
            for x in rows:
                period = "-".join(filter(None, [x.get('start_date'), x.get('end_date')]))
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                r = p.add_run(f"{x.get('organization')} - {x.get('role')}: {x.get('course_name')} ({period})")
                r.bold = True
                if x.get('description'):
                    add_complete_text(doc, x['description'], bullet=True)

    if options.get('organizations', True):
        rows = sections.get('organizations', [])
        if rows:
            add_section_heading(doc, "Professional Organizations")
            for x in rows:
                years = "-".join(filter(None, [x.get('start_year'), x.get('end_year')]))
                role = f" - {x.get('role')}" if x.get('role') else ''
                add_bullet(doc, f"{years} {x.get('organization')}{role}")

    if options.get('certifications', True):
        rows = sections.get('certifications', [])
        if rows:
            add_section_heading(doc, "Digital Forensic Certifications")
            for x in rows:
                earned = pretty_date(full_text(x.get('earned_date')))
                expires = f" (expires {pretty_date(full_text(x.get('expiration_date')))})" if x.get('expiration_date') else ''
                add_bullet(doc, f"{full_text(x.get('certification'))}, {full_text(x.get('issuing_organization'))} {earned}{expires}".strip())

    if options.get('skills', True):
        rows = [item for values in sections.get('skills', {}).values() for item in values]
        if rows:
            add_section_heading(doc, "Relevant Skills and Tools")
            by_cat: dict[str, list[str]] = {}
            for x in rows:
                by_cat.setdefault(x.get('category') or 'Other', []).append(x.get('skill', ''))
            for cat, values in by_cat.items():
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                r = p.add_run(f"{cat}: ")
                r.bold = True
                p.add_run('; '.join(values))

    if options.get('education', True):
        rows = sections.get('education', [])
        if rows:
            add_section_heading(doc, "Education")
            for x in rows:
                honors = f" ({x.get('honors')})" if x.get('honors') else ''
                add_bullet(doc, f"{x.get('degree')} - {x.get('institution')}, {x.get('graduation_date')}{honors}")

    if options.get('testimony', True):
        testimony_groups = sections.get('testimony', {})
        rows = testimony_groups.get('Expert Witness', []) + testimony_groups.get('Fact Witness', [])
        if rows:
            add_section_heading(doc, "Courtroom Testimony")
            for kind in ('Expert Witness', 'Fact Witness'):
                subset = [x for x in rows if x.get('witness_type') == kind]
                if not subset:
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                r = p.add_run(kind)
                r.bold = True
                for x in subset:
                    add_bullet(doc, f"SA# {x.get('case_number')}, {x.get('court')} {x.get('jurisdiction')}, {pretty_date(x.get('testimony_date'))}")

    if options.get('full_training', False):
        training_section = sections.get('full_training')
        rows = training_section.get('rows', []) if training_section else []
        if rows:
            doc.add_section(WD_SECTION.NEW_PAGE)
            add_section_heading(doc, "Training Courses")
            total = training_section.get('total_hours', 0)
            p = doc.add_paragraph(f"Documented training hours in database: {total:,.2f}")
            p.runs[0].bold = True
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            headers = ['Date', 'Course', 'Provider', 'Hours']
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                set_cell_shading(table.rows[0].cells[i], 'D9EAF7')
                for run in table.rows[0].cells[i].paragraphs[0].runs:
                    run.bold = True
            for x in rows:
                cells = table.add_row().cells
                vals = [full_text(x.get('attended_date')), full_text(x.get('course_name')), full_text(x.get('provider')), '' if x.get('hours') is None else full_text(x.get('hours'))]
                for i, val in enumerate(vals):
                    cells[i].text = val

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run(f"Generated {date.today().strftime('%B %d, %Y')}").font.size = Pt(8)

    doc.save(output_path)
    return output_path
