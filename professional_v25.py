from __future__ import annotations

"""
Forensic CV Manager Professional Edition v2.5 productivity/records upgrade.

This module layers on top of the Professional Tracking extension and adds:
- Global cross-table search
- Per-tab filters and record counts
- Dashboard drill-down
- Annual activity statistics/reports
- Canonical forensic tool library + Case Work multi-select
- Case Index + linked-record explorer
- Export presets
- Expert qualification / voir-dire report
- Certification/CPE alerts
- Stronger input validation / duplicate warnings
- Automatic daily backups with 30-backup retention
- Database schema versioning
- Professional About / Changelog views

It intentionally avoids storing or exporting additional investigative narrative
unless the user explicitly enters/selects it.
"""

import csv
import html
import json
import re
import shutil
import sys
import tempfile
from collections import Counter, defaultdict
from datetime import date, datetime
from pathlib import Path
from typing import Any, Iterable

import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk

import database as database_module
from date_utils import date_sort_key, normalize_date

try:
    import professional_tracking as pt
except Exception:
    pt = None


PRO_EDITION = "Professional Edition"
PRO_VERSION = "2.5.0"
SCHEMA_VERSION = 5
BACKUP_RETENTION = 30

CHANGELOG = """FORENSIC CV MANAGER PROFESSIONAL — VERSION 2.5.0

Navigation and usability
• Added global cross-table search.
• Added field filters and year filters to record tabs.
• Added dashboard metric drill-down.
• Added categorized navigation compatibility.

Forensic case experience
• Added Case Index for linking professional records by case number.
• Added linked-record Case Explorer.
• Added canonical Tool Library.
• Added multi-select forensic tools for Case Work.
• Added annual casework statistics and tool/device/acquisition breakdowns.

Reporting
• Added export presets.
• Added Expert Qualification / Voir Dire report.
• Added Annual Professional Activity report.
• Added granular detailed appendix/privacy controls from the Professional Tracking upgrade.

Credential management
• Added certification expiration and CPE/CE progress alerts.
• Added stronger duplicate/date/numeric/storage-size validation.

Data protection and maintenance
• Added automatic daily SQLite backups.
• Keeps the newest 30 automatic backups.
• Added explicit database schema version metadata and migration support.

Release information
• Professional Edition version set to 2.5.0.
• Added expanded About information and changelog view.
"""


# ---------------------------------------------------------------------------
# Database schema / migration
# ---------------------------------------------------------------------------

CASE_INDEX_FIELDS = [
    "case_number", "opened_date", "closed_date", "agency", "case_type",
    "status", "notes", "sort_order",
]

TOOL_LIBRARY_FIELDS = [
    "canonical_name", "category", "aliases", "active", "sort_order",
]

NEW_TABLES = {
    "case_index": CASE_INDEX_FIELDS,
    "tool_library": TOOL_LIBRARY_FIELDS,
}

LINKED_TABLES = (
    "casework",
    "forensic_reports",
    "peer_reviews",
    "testimony",
    "court_qualifications",
)

ORDER_BY = {
    "case_index": "opened_date DESC, case_number COLLATE NOCASE, sort_order, id DESC",
    "tool_library": "category COLLATE NOCASE, canonical_name COLLATE NOCASE, sort_order, id DESC",
}

DEFAULT_TOOLS = [
    ("Magnet AXIOM", "Computer / Mobile Forensics", "AXIOM; Magnet Axiom"),
    ("Cellebrite Physical Analyzer", "Mobile Forensics", "Cellebrite PA; Physical Analyzer"),
    ("Cellebrite UFED", "Mobile Forensics", "UFED; Cellebrite UFED"),
    ("X-Ways Forensics", "Computer Forensics", "X-Ways; XWays; X-Ways"),
    ("EnCase Forensic", "Computer Forensics", "EnCase"),
    ("FTK", "Computer Forensics", "AccessData FTK; Exterro FTK"),
    ("GrayKey", "Mobile Forensics", "GrayKey; Grayshift GrayKey"),
    ("Autopsy", "Computer Forensics", "Sleuth Kit Autopsy"),
    ("Volatility", "Memory Forensics", "Volatility Framework"),
    ("Wireshark", "Network Forensics", "Wireshark"),
    ("Ghidra", "Malware / Reverse Engineering", "NSA Ghidra"),
    ("IDA Pro", "Malware / Reverse Engineering", "IDA; IDA Professional"),
]


def _sql_type(field: str) -> str:
    if field in {"active"}:
        return "INTEGER DEFAULT 1"
    if field == "sort_order":
        return "INTEGER DEFAULT 0"
    return "TEXT"


def _create_table_sql(table: str, fields: Iterable[str]) -> str:
    cols = [
        "id INTEGER PRIMARY KEY AUTOINCREMENT",
        "profile_id INTEGER NOT NULL",
    ]
    cols.extend(f"{field} {_sql_type(field)}" for field in fields)
    cols.append("FOREIGN KEY(profile_id) REFERENCES profiles(id) ON DELETE CASCADE")
    return f"CREATE TABLE IF NOT EXISTS {table} ({', '.join(cols)})"


_DB_INSTALLED = False


def install_database_extensions() -> None:
    """Install schema versioning, Case Index, Tool Library, and linkage support."""
    global _DB_INSTALLED
    if _DB_INSTALLED:
        return

    for table, fields in NEW_TABLES.items():
        if table not in database_module.DATA_TABLES:
            database_module.DATA_TABLES.append(table)
        database_module.TABLE_FIELDS.setdefault(table, list(fields))
        database_module.CREATE_DATA_TABLES.setdefault(
            table, _create_table_sql(table, database_module.TABLE_FIELDS[table])
        )

    # case_id is maintained internally, not shown in the generic form.
    for table in LINKED_TABLES:
        if table in database_module.TABLE_FIELDS and "case_id" not in database_module.TABLE_FIELDS[table]:
            database_module.TABLE_FIELDS[table].append("case_id")

    Database = database_module.Database
    old_migrate = Database._migrate_schema
    old_init = Database.__init__
    old_list_rows = Database.list_rows
    old_insert = Database.insert_row
    old_update = Database.update_row
    old_delete = Database.delete_row
    old_create_profile = getattr(Database, "create_profile", None)

    def migrate(self):
        old_migrate(self)

        self.conn.execute(
            "CREATE TABLE IF NOT EXISTS schema_meta "
            "(key TEXT PRIMARY KEY, value TEXT NOT NULL)"
        )

        for table in LINKED_TABLES:
            if self._table_exists(table) and "case_id" not in self._columns(table):
                self.conn.execute(f"ALTER TABLE {table} ADD COLUMN case_id INTEGER")

        self.conn.execute(
            "CREATE UNIQUE INDEX IF NOT EXISTS idx_case_index_profile_number "
            "ON case_index(profile_id, case_number COLLATE NOCASE)"
        )
        self.conn.execute(
            "CREATE UNIQUE INDEX IF NOT EXISTS idx_tool_library_profile_name "
            "ON tool_library(profile_id, canonical_name COLLATE NOCASE)"
        )
        self.conn.execute(
            "INSERT OR REPLACE INTO schema_meta(key,value) VALUES('schema_version',?)",
            (str(SCHEMA_VERSION),),
        )
        self.conn.execute(
            "INSERT OR REPLACE INTO schema_meta(key,value) VALUES('professional_version',?)",
            (PRO_VERSION,),
        )
        self.conn.commit()

    def init(self, path):
        old_init(self, path)
        seed_default_tools(self)
        sync_case_index(self)

    def list_rows(self, table: str, order_by: str | None = None):
        if order_by is None and table in ORDER_BY:
            order_by = ORDER_BY[table]
        return old_list_rows(self, table, order_by)

    def insert_row(self, table: str, data: dict[str, Any]) -> int:
        row_id = old_insert(self, table, data)
        if table in LINKED_TABLES:
            link_row_to_case(self, table, row_id)
        elif table == "case_index":
            sync_case_index(self)
        return row_id

    def update_row(self, table: str, row_id: int, data: dict[str, Any]) -> None:
        old_update(self, table, row_id, data)
        if table in LINKED_TABLES:
            link_row_to_case(self, table, row_id)
        elif table == "case_index":
            sync_case_index(self)

    def delete_row(self, table: str, row_id: int) -> None:
        old_delete(self, table, row_id)
        if table == "case_index":
            # Child rows remain valid; case_id is cleared if the index record is gone.
            for child in LINKED_TABLES:
                if self._table_exists(child) and "case_id" in self._columns(child):
                    self.conn.execute(
                        f"UPDATE {child} SET case_id=NULL "
                        "WHERE case_id NOT IN (SELECT id FROM case_index)"
                    )
            self.conn.commit()

    def create_profile(self, profile_name: str, full_name: str = "") -> int:
        if old_create_profile is None:
            raise AttributeError("Database.create_profile is unavailable")
        pid = old_create_profile(self, profile_name, full_name)
        seed_default_tools_for_profile(self, pid)
        return pid

    Database._migrate_schema = migrate
    Database.__init__ = init
    Database.list_rows = list_rows
    Database.insert_row = insert_row
    Database.update_row = update_row
    Database.delete_row = delete_row
    if old_create_profile is not None:
        Database.create_profile = create_profile
    _DB_INSTALLED = True


def get_schema_version(db) -> int:
    try:
        row = db.conn.execute(
            "SELECT value FROM schema_meta WHERE key='schema_version'"
        ).fetchone()
        return int(row[0]) if row else 0
    except Exception:
        return 0


def seed_default_tools_for_profile(db, profile_id: int) -> None:
    try:
        pid = int(profile_id)
        count = db.conn.execute(
            "SELECT COUNT(*) FROM tool_library WHERE profile_id=?",
            (pid,),
        ).fetchone()[0]
        if count:
            return
        for name, category, aliases in DEFAULT_TOOLS:
            db.conn.execute(
                "INSERT INTO tool_library "
                "(profile_id,canonical_name,category,aliases,active,sort_order) "
                "VALUES (?,?,?,?,1,0)",
                (pid, name, category, aliases),
            )
        db.conn.commit()
    except Exception:
        pass


def seed_default_tools(db) -> None:
    try:
        for profile in db.list_profiles():
            seed_default_tools_for_profile(db, int(profile["id"]))
    except Exception:
        pass


def _case_number(value: Any) -> str:
    return str(value or "").strip()


def ensure_case_index_row(db, profile_id: int, case_number: str, source: dict[str, Any] | None = None) -> int | None:
    number = _case_number(case_number)
    if not number:
        return None
    row = db.conn.execute(
        "SELECT id FROM case_index WHERE profile_id=? AND case_number=? COLLATE NOCASE",
        (profile_id, number),
    ).fetchone()
    if row:
        case_id = int(row[0])
    else:
        source = source or {}
        agency = source.get("requesting_agency") or source.get("agency") or ""
        case_type = source.get("case_type") or ""
        status = source.get("status") or ""
        opened = (
            source.get("examination_date")
            or source.get("report_date")
            or source.get("review_date")
            or source.get("testimony_date")
            or source.get("qualification_date")
            or ""
        )
        cur = db.conn.execute(
            "INSERT INTO case_index "
            "(profile_id,case_number,opened_date,agency,case_type,status,sort_order) "
            "VALUES (?,?,?,?,?,?,0)",
            (profile_id, number, opened, agency, case_type, status),
        )
        case_id = int(cur.lastrowid)
    return case_id


def link_row_to_case(db, table: str, row_id: int) -> None:
    if table not in LINKED_TABLES or not db._table_exists(table):
        return
    row = db.conn.execute(
        f"SELECT * FROM {table} WHERE id=?",
        (row_id,),
    ).fetchone()
    if not row:
        return
    data = dict(row)
    number = _case_number(data.get("case_number"))
    if not number:
        if "case_id" in db._columns(table):
            db.conn.execute(f"UPDATE {table} SET case_id=NULL WHERE id=?", (row_id,))
            db.conn.commit()
        return
    pid = int(data.get("profile_id") or db.current_profile_id)
    case_id = ensure_case_index_row(db, pid, number, data)
    if case_id is not None and "case_id" in db._columns(table):
        db.conn.execute(f"UPDATE {table} SET case_id=? WHERE id=?", (case_id, row_id))
        db.conn.commit()


def sync_case_index(db) -> None:
    """Create Case Index entries from existing records and populate case_id."""
    try:
        for table in LINKED_TABLES:
            if not db._table_exists(table) or "case_number" not in db._columns(table):
                continue
            rows = db.conn.execute(
                f"SELECT * FROM {table} WHERE case_number IS NOT NULL "
                "AND TRIM(case_number)<>''"
            ).fetchall()
            for raw in rows:
                data = dict(raw)
                pid = int(data.get("profile_id") or 0)
                if not pid:
                    continue
                case_id = ensure_case_index_row(db, pid, data.get("case_number"), data)
                if case_id and "case_id" in db._columns(table):
                    db.conn.execute(
                        f"UPDATE {table} SET case_id=? WHERE id=?",
                        (case_id, data["id"]),
                    )
        db.conn.commit()
    except Exception:
        pass


# ---------------------------------------------------------------------------
# TABLE_CONFIG extension
# ---------------------------------------------------------------------------

V25_TABLE_CONFIG = {
    "case_index": {
        "label": "Case Index",
        "display": ["case_number", "opened_date", "agency", "case_type", "status"],
        "fields": [
            ("case_number", "Case Number / Reference", "entry"),
            ("opened_date", "Opened / First Activity Date", "entry"),
            ("closed_date", "Closed Date", "entry"),
            ("agency", "Agency / Organization", "entry"),
            ("case_type", "Case Type", "entry"),
            ("status", "Status", "combo", ["", "Open", "Ongoing", "Complete", "Closed", "Archived"]),
            ("notes", "General Professional Notes", "text"),
        ],
    },
    "tool_library": {
        "label": "Tool Library",
        "display": ["canonical_name", "category", "aliases", "active"],
        "fields": [
            ("canonical_name", "Canonical Tool Name", "entry"),
            ("category", "Category", "entry"),
            ("aliases", "Aliases (semicolon separated)", "entry"),
            ("active", "Active / Show in Tool Selector", "check"),
        ],
    },
}

V25_DATE_FIELDS = {"opened_date", "closed_date"}

FILTER_CONFIG = {
    "case_index": ["case_type", "status", "agency"],
    "casework": ["case_type", "device_type", "acquisition_type", "status"],
    "forensic_reports": ["report_type", "role", "subject_area"],
    "peer_reviews": ["review_type", "role", "outcome"],
    "court_qualifications": ["jurisdiction", "qualification_area", "ruling"],
    "testimony": ["witness_type", "proceeding_type", "jurisdiction"],
    "tool_experience": ["category", "proficiency"],
    "tool_library": ["category", "active"],
    "training": ["category", "provider", "delivery_method"],
    "certifications": ["status", "issuing_organization"],
    "teaching": ["organization", "role", "curriculum_role"],
    "presentations": ["organization", "role"],
    "publications": ["publication_type", "authorship_role"],
    "validations": ["validation_type", "result", "tool_name"],
    "procedures": ["document_type", "role", "status"],
    "mentoring": ["role", "program"],
    "projects": ["project_type", "organization"],
    "professional_evidence": ["evidence_type", "related_area"],
    "employment": ["employer", "title"],
    "education": ["institution", "degree"],
    "organizations": ["organization", "role"],
    "skills": ["category", "proficiency"],
    "achievements": ["category", "organization"],
}

DATE_FIELD_PRIORITY = [
    "examination_date", "report_date", "review_date", "qualification_date",
    "testimony_date", "attended_date", "earned_date", "presentation_date",
    "publication_date", "validation_date", "effective_date", "evidence_date",
    "achievement_date", "start_date", "graduation_date", "opened_date",
]


def extend_table_config(table_config, date_fields, year_fields=None) -> None:
    date_fields.update(V25_DATE_FIELDS)
    for table, config in V25_TABLE_CONFIG.items():
        table_config.setdefault(
            table,
            {
                "label": config["label"],
                "display": list(config["display"]),
                "fields": list(config["fields"]),
            },
        )


# ---------------------------------------------------------------------------
# Common helpers
# ---------------------------------------------------------------------------

def _root_app(widget):
    root = widget.winfo_toplevel()
    return root


def _rows(db, table: str) -> list[dict[str, Any]]:
    try:
        return db.list_rows(table)
    except Exception:
        return []


def _to_float(value: Any) -> float:
    try:
        return float(value or 0)
    except (ValueError, TypeError):
        return 0.0


def _to_int(value: Any) -> int:
    try:
        return int(float(value or 0))
    except (ValueError, TypeError):
        return 0


def _truthy(value: Any) -> bool:
    if isinstance(value, str):
        return value.strip().lower() in {"1", "yes", "true", "y", "active"}
    return bool(value)


def _field_year(value: Any) -> str:
    text = str(value or "").strip()
    m = re.match(r"^(\d{4})", text)
    return m.group(1) if m else ""


_SIZE_RE = re.compile(
    r"^\s*([\d,.]+)\s*(B|KB|MB|GB|TB|PB|KIB|MIB|GIB|TIB|PIB)?\s*$",
    re.I,
)


def parse_size_bytes(value: Any) -> float | None:
    text = str(value or "").strip()
    if not text:
        return 0.0
    text = text.replace(",", "")
    m = _SIZE_RE.match(text)
    if not m:
        return None
    number = float(m.group(1))
    unit = (m.group(2) or "B").upper()
    powers = {
        "B": 0, "KB": 1, "KIB": 1,
        "MB": 2, "MIB": 2,
        "GB": 3, "GIB": 3,
        "TB": 4, "TIB": 4,
        "PB": 5, "PIB": 5,
    }
    return number * (1024 ** powers[unit])


def format_bytes(value: float) -> str:
    units = ["B", "KB", "MB", "GB", "TB", "PB"]
    size = float(value or 0)
    for unit in units:
        if size < 1024 or unit == units[-1]:
            return f"{size:,.2f} {unit}"
        size /= 1024
    return f"{size:,.2f} PB"


def _tool_library(db) -> list[dict[str, Any]]:
    try:
        return db.list_rows("tool_library")
    except Exception:
        return []


def _tool_alias_map(db) -> dict[str, str]:
    mapping: dict[str, str] = {}
    for row in _tool_library(db):
        canonical = str(row.get("canonical_name") or "").strip()
        if not canonical:
            continue
        mapping[canonical.casefold()] = canonical
        for alias in re.split(r"[;\n]+", str(row.get("aliases") or "")):
            alias = alias.strip()
            if alias:
                mapping[alias.casefold()] = canonical
    return mapping


def normalize_tools(db, value: Any) -> list[str]:
    mapping = _tool_alias_map(db)
    output = []
    seen = set()
    for raw in re.split(r"[;\n]+", str(value or "")):
        name = raw.strip()
        if not name:
            continue
        canonical = mapping.get(name.casefold(), name)
        key = canonical.casefold()
        if key not in seen:
            seen.add(key)
            output.append(canonical)
    return output


def canonical_tool_counter(db, rows: Iterable[dict[str, Any]]) -> Counter:
    counter = Counter()
    for row in rows:
        counter.update(normalize_tools(db, row.get("tools_used")))
    return counter


def _first_date_field(table: str, row_or_fields: Any) -> str | None:
    if isinstance(row_or_fields, dict):
        fields = row_or_fields.keys()
    else:
        fields = row_or_fields
    fields = set(fields)
    for field in DATE_FIELD_PRIORITY:
        if field in fields:
            return field
    return None


# ---------------------------------------------------------------------------
# Automatic daily backups
# ---------------------------------------------------------------------------

_BACKUP_WRAPPED = False


def install_auto_backup(app_module) -> None:
    global _BACKUP_WRAPPED
    if _BACKUP_WRAPPED:
        return
    original = app_module.prepare_portable_database

    def wrapped():
        db_path = Path(original())
        try:
            if db_path.exists() and db_path.stat().st_size > 0:
                backup_dir = db_path.parent.parent / "Backups" / "Auto"
                backup_dir.mkdir(parents=True, exist_ok=True)
                today = date.today().isoformat()
                target = backup_dir / f"forensic_cv_auto_{today}.sqlite3"
                if not target.exists():
                    shutil.copy2(db_path, target)
                backups = sorted(
                    backup_dir.glob("forensic_cv_auto_*.sqlite3"),
                    key=lambda p: p.stat().st_mtime,
                    reverse=True,
                )
                for old in backups[BACKUP_RETENTION:]:
                    try:
                        old.unlink()
                    except OSError:
                        pass
        except Exception:
            # Backup failure must never prevent the application from opening.
            pass
        return db_path

    app_module.prepare_portable_database = wrapped
    _BACKUP_WRAPPED = True


# ---------------------------------------------------------------------------
# Enhanced Record Dialog + validation + tool multi-select
# ---------------------------------------------------------------------------

FLOAT_FIELDS = {
    "hours", "training_hours", "cpe_credits", "cpe_required", "cpe_earned",
    "testimony_hours",
}
INT_FIELDS = {
    "pages", "audience_size", "approx_examinations", "students",
}
BOOL_FIELDS = {
    "active", "supplemental", "peer_reviewed", "changes_required",
    "report_written", "testified", "qualified_expert", "voir_dire",
    "core_training",
}

ALL_DATE_FIELDS = set(V25_DATE_FIELDS)
if pt:
    ALL_DATE_FIELDS.update(getattr(pt, "EXTENDED_DATE_FIELDS", set()))
ALL_DATE_FIELDS.update({
    "start_date", "end_date", "graduation_date", "attended_date",
    "expiration_date", "earned_date", "testimony_date", "achievement_date",
})


class EnhancedRecordDialog(tk.Toplevel):
    def __init__(self, parent, table: str, config: dict[str, Any], initial=None):
        super().__init__(parent)
        self.table = table
        self.config_data = config
        self.initial = dict(initial or {})
        self.initial_id = self.initial.get("id")
        self.db = getattr(parent, "db", None)
        self.result = None
        self.vars: dict[str, Any] = {}
        self.widgets: dict[str, Any] = {}
        self.tool_listbox = None
        self.other_tools_var = None

        self.title(("Edit " if initial else "Add ") + config["label"])
        self.transient(parent)
        self.grab_set()
        self.resizable(True, True)

        outer = ttk.Frame(self)
        outer.pack(fill="both", expand=True)
        canvas = tk.Canvas(outer, highlightthickness=0)
        scroll = ttk.Scrollbar(outer, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=scroll.set)
        scroll.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)

        form = ttk.Frame(canvas, padding=12)
        window = canvas.create_window((0, 0), window=form, anchor="nw")
        form.columnconfigure(1, weight=1)

        def sync_region(_=None):
            canvas.configure(scrollregion=canvas.bbox("all"))

        def sync_width(event):
            canvas.itemconfigure(window, width=event.width)

        form.bind("<Configure>", sync_region)
        canvas.bind("<Configure>", sync_width)

        row_index = 0
        for field in config["fields"]:
            name, label, kind, *extra = field
            ttk.Label(form, text=label + ":").grid(
                row=row_index, column=0, sticky="nw", padx=8, pady=5
            )
            value = "" if initial is None or initial.get(name) is None else initial.get(name)

            if table == "casework" and name == "tools_used":
                container = ttk.Frame(form)
                container.grid(row=row_index, column=1, sticky="nsew", padx=8, pady=5)
                container.columnconfigure(0, weight=1)
                ttk.Label(
                    container,
                    text="Select canonical tools (Ctrl-click for multiple):"
                ).grid(row=0, column=0, sticky="w")
                lb = tk.Listbox(
                    container, selectmode="multiple", exportselection=False, height=8
                )
                lb.grid(row=1, column=0, sticky="nsew", pady=(3, 6))
                tools = [
                    r for r in _tool_library(self.db)
                    if _truthy(r.get("active", 1))
                ] if self.db else []
                self._tool_names = [
                    str(r.get("canonical_name") or "").strip()
                    for r in tools
                    if str(r.get("canonical_name") or "").strip()
                ]
                existing = normalize_tools(self.db, value) if self.db else []
                existing_cf = {x.casefold() for x in existing}
                library_cf = {x.casefold() for x in self._tool_names}
                for idx, tool in enumerate(self._tool_names):
                    lb.insert("end", tool)
                    if tool.casefold() in existing_cf:
                        lb.selection_set(idx)
                extras = [x for x in existing if x.casefold() not in library_cf]
                ttk.Label(container, text="Additional / unlisted tools:").grid(
                    row=2, column=0, sticky="w"
                )
                self.other_tools_var = tk.StringVar(value="; ".join(extras))
                ttk.Entry(container, textvariable=self.other_tools_var).grid(
                    row=3, column=0, sticky="ew", pady=(3, 0)
                )
                self.tool_listbox = lb
                self.widgets[name] = container
            elif kind == "text":
                widget = tk.Text(form, width=65, height=4, wrap="word")
                widget.insert("1.0", str(value))
                widget.grid(row=row_index, column=1, sticky="nsew", padx=8, pady=5)
                self.widgets[name] = widget
            elif kind == "check":
                var = tk.IntVar(value=int(value or 0))
                widget = ttk.Checkbutton(form, variable=var)
                widget.grid(row=row_index, column=1, sticky="w", padx=8, pady=5)
                self.vars[name] = var
                self.widgets[name] = widget
            elif kind == "combo":
                var = tk.StringVar(value=str(value))
                vals = extra[0] if extra else []
                widget = ttk.Combobox(
                    form, textvariable=var, values=vals, state="readonly"
                )
                if not value and vals:
                    var.set(vals[0])
                widget.grid(row=row_index, column=1, sticky="ew", padx=8, pady=5)
                self.vars[name] = var
                self.widgets[name] = widget
            else:
                var = tk.StringVar(value=str(value))
                widget = ttk.Entry(form, textvariable=var, width=65)
                widget.grid(row=row_index, column=1, sticky="ew", padx=8, pady=5)
                self.vars[name] = var
                self.widgets[name] = widget
            row_index += 1

        buttons = ttk.Frame(form)
        buttons.grid(row=row_index, column=0, columnspan=2, sticky="e", padx=8, pady=10)
        ttk.Button(buttons, text="Cancel", command=self.destroy).pack(side="right", padx=4)
        ttk.Button(buttons, text="Save", command=self.save).pack(side="right", padx=4)

        self.geometry(f"900x{min(800, max(620, self.winfo_screenheight() - 130))}")
        self.bind("<Escape>", lambda e: self.destroy())

        def wheel(event):
            canvas.yview_scroll((-1 if event.delta > 0 else 1) * 3, "units")

        canvas.bind_all("<MouseWheel>", wheel)
        self._wheel_bound = True

    def destroy(self):
        if getattr(self, "_wheel_bound", False):
            try:
                self.unbind_all("<MouseWheel>")
            except Exception:
                pass
        super().destroy()

    def _collect_value(self, field: str):
        if field == "tools_used" and self.tool_listbox is not None:
            selected = [
                self._tool_names[i]
                for i in self.tool_listbox.curselection()
            ]
            extras = normalize_tools(
                self.db,
                self.other_tools_var.get() if self.other_tools_var else ""
            )
            combined = []
            seen = set()
            for name in selected + extras:
                key = name.casefold()
                if key not in seen:
                    seen.add(key)
                    combined.append(name)
            return "; ".join(combined)

        widget = self.widgets[field]
        if isinstance(widget, tk.Text):
            return widget.get("1.0", "end").strip()
        return self.vars[field].get()

    def _validate_duplicates(self, result: dict[str, Any]) -> bool:
        if not self.db:
            return True

        if self.table == "case_index":
            number = _case_number(result.get("case_number"))
            if number:
                for row in _rows(self.db, "case_index"):
                    if row.get("id") == self.initial_id:
                        continue
                    if _case_number(row.get("case_number")).casefold() == number.casefold():
                        messagebox.showerror(
                            "Duplicate Case",
                            "That case number already exists in Case Index. "
                            "Edit the existing Case Index record instead.",
                            parent=self,
                        )
                        return False

        if self.table == "tool_library":
            name = str(result.get("canonical_name") or "").strip()
            if name:
                for row in _rows(self.db, "tool_library"):
                    if row.get("id") == self.initial_id:
                        continue
                    if str(row.get("canonical_name") or "").strip().casefold() == name.casefold():
                        messagebox.showerror(
                            "Duplicate Tool",
                            "That canonical tool name already exists in Tool Library.",
                            parent=self,
                        )
                        return False

        if self.table == "certifications":
            cert = str(result.get("certification") or "").strip()
            if cert:
                for row in _rows(self.db, "certifications"):
                    if row.get("id") == self.initial_id:
                        continue
                    if str(row.get("certification") or "").strip().casefold() == cert.casefold():
                        return messagebox.askyesno(
                            "Possible Duplicate Certification",
                            f'A certification named "{cert}" already exists. Save another record anyway?',
                            parent=self,
                        )

        if self.table == "casework":
            case_no = _case_number(result.get("case_number"))
            evidence = str(result.get("evidence_number") or "").strip()
            device = str(result.get("device_make_model") or "").strip()
            if case_no:
                for row in _rows(self.db, "casework"):
                    if row.get("id") == self.initial_id:
                        continue
                    same_case = _case_number(row.get("case_number")).casefold() == case_no.casefold()
                    same_item = evidence and str(row.get("evidence_number") or "").strip().casefold() == evidence.casefold()
                    same_device = (not evidence and device and
                                   str(row.get("device_make_model") or "").strip().casefold() == device.casefold())
                    if same_case and (same_item or same_device):
                        return messagebox.askyesno(
                            "Possible Duplicate Examination",
                            "A Case Work record appears to use the same case and "
                            "evidence/device reference. Save it anyway?",
                            parent=self,
                        )
        return True

    def save(self):
        result = {}
        for field in self.widgets:
            value = self._collect_value(field)

            if field in FLOAT_FIELDS:
                try:
                    value = float(value) if str(value).strip() else None
                except ValueError:
                    messagebox.showerror(
                        "Invalid Number",
                        f"{field.replace('_', ' ').title()} must be a number.",
                        parent=self,
                    )
                    return
                if value is not None and value < 0:
                    messagebox.showerror(
                        "Invalid Number", "Negative values are not allowed.", parent=self
                    )
                    return
            elif field in INT_FIELDS:
                try:
                    value = int(float(value)) if str(value).strip() else None
                except ValueError:
                    messagebox.showerror(
                        "Invalid Number",
                        f"{field.replace('_', ' ').title()} must be a whole number.",
                        parent=self,
                    )
                    return
                if value is not None and value < 0:
                    messagebox.showerror(
                        "Invalid Number", "Negative values are not allowed.", parent=self
                    )
                    return
            elif field in ALL_DATE_FIELDS:
                try:
                    value = normalize_date(
                        value,
                        allow_present=field in {"end_date", "last_used"},
                        allow_year_only=True,
                    )
                except ValueError as exc:
                    messagebox.showerror("Invalid Date", str(exc), parent=self)
                    return
            elif field in {"start_year", "end_year"}:
                try:
                    value = normalize_date(
                        value,
                        allow_present=field == "end_year",
                        allow_year_only=True,
                    )
                    if value not in ("", "Present") and len(value) != 4:
                        raise ValueError("Enter a four-digit year.")
                except ValueError as exc:
                    messagebox.showerror("Invalid Year", str(exc), parent=self)
                    return

            if field in {"data_acquired", "data_examined"} and str(value or "").strip():
                parsed = parse_size_bytes(value)
                if parsed is None:
                    messagebox.showerror(
                        "Invalid Data Size",
                        f'Enter {field.replace("_", " ")} like "938 GB", "1.5 TB", or "240 MB".',
                        parent=self,
                    )
                    return

            result[field] = value

        # Date range logic.
        for start_field, end_field in (
            ("start_date", "end_date"),
            ("first_used", "last_used"),
            ("opened_date", "closed_date"),
        ):
            start = result.get(start_field)
            end = result.get(end_field)
            if start and end and str(end) != "Present":
                if date_sort_key(start) > date_sort_key(end):
                    messagebox.showerror(
                        "Invalid Date Range",
                        f"{end_field.replace('_', ' ').title()} cannot be before "
                        f"{start_field.replace('_', ' ').title()}.",
                        parent=self,
                    )
                    return

        if not self._validate_duplicates(result):
            return

        self.result = result
        self.destroy()


# ---------------------------------------------------------------------------
# Enhanced per-tab search/filter UI
# ---------------------------------------------------------------------------

class EnhancedRecordsTab(ttk.Frame):
    def __init__(self, parent, db, table: str, status_callback):
        super().__init__(parent)
        app_module = sys.modules.get(parent.winfo_toplevel().__class__.__module__)
        self.db = db
        self.table = table
        self.config_data = app_module.TABLE_CONFIG[table]
        self.status_callback = status_callback
        self.search_var = tk.StringVar()
        self.sort_column = None
        self.sort_descending = False
        self.filter_vars: dict[str, tk.StringVar] = {}
        self.year_var = tk.StringVar(value="All")

        top = ttk.Frame(self)
        top.pack(fill="x", padx=8, pady=(8, 4))
        ttk.Label(top, text="Search:").pack(side="left")
        search = ttk.Entry(top, textvariable=self.search_var, width=34)
        search.pack(side="left", padx=5)
        self.search_var.trace_add("write", lambda *_: self.refresh())

        ttk.Button(top, text="Add", command=self.add).pack(side="right", padx=3)
        ttk.Button(top, text="Edit", command=self.edit).pack(side="right", padx=3)
        ttk.Button(top, text="Delete", command=self.delete).pack(side="right", padx=3)

        if table == "case_index":
            ttk.Button(
                top, text="View Linked Records", command=self.view_linked_records
            ).pack(side="right", padx=3)
        if table == "tool_library":
            ttk.Button(
                top, text="Normalize Existing Case Work", command=self.normalize_casework_tools
            ).pack(side="right", padx=3)

        self.filter_frame = ttk.LabelFrame(self, text="Filters", padding=(6, 4))
        self.filter_frame.pack(fill="x", padx=8, pady=(0, 6))
        self._build_filters()

        content = ttk.Frame(self)
        content.pack(fill="both", expand=True, padx=8, pady=(0, 8))
        content.rowconfigure(0, weight=1)
        content.columnconfigure(0, weight=1)

        cols = self.config_data["display"]
        self.tree = ttk.Treeview(
            content, columns=cols, show="headings", selectmode="browse"
        )
        for col in cols:
            self.tree.heading(
                col,
                text=self._heading_text(col),
                command=lambda column=col: self.sort_by(column),
            )
            width = 120
            if col in {
                "course_name", "certification", "employer", "degree", "skill",
                "achievement", "case_number", "canonical_name", "title",
                "device_make_model",
            }:
                width = 220
            self.tree.column(col, width=width, minwidth=70, stretch=True)

        y = ttk.Scrollbar(content, orient="vertical", command=self.tree.yview)
        x = ttk.Scrollbar(content, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=y.set, xscrollcommand=x.set)
        self.tree.grid(row=0, column=0, sticky="nsew")
        y.grid(row=0, column=1, sticky="ns")
        x.grid(row=1, column=0, sticky="ew")
        self.tree.bind("<Double-1>", lambda e: self.edit())

        self.count_label = ttk.Label(self, text="")
        self.count_label.pack(anchor="e", padx=10, pady=(0, 5))
        self.refresh()

    def _build_filters(self):
        for child in self.filter_frame.winfo_children():
            child.destroy()
        fields = FILTER_CONFIG.get(self.table, [])[:4]
        rows = _rows(self.db, self.table)

        col = 0
        for field in fields:
            if not rows or field not in rows[0]:
                continue
            ttk.Label(
                self.filter_frame,
                text=field.replace("_", " ").title() + ":"
            ).grid(row=0, column=col, sticky="w", padx=(0, 3))
            var = tk.StringVar(value="All")
            values = sorted({
                str(r.get(field) if r.get(field) is not None else "").strip()
                for r in rows
                if str(r.get(field) if r.get(field) is not None else "").strip()
            }, key=str.casefold)
            combo = ttk.Combobox(
                self.filter_frame,
                textvariable=var,
                values=["All"] + values,
                state="readonly",
                width=18,
            )
            combo.grid(row=0, column=col + 1, sticky="w", padx=(0, 8))
            combo.bind("<<ComboboxSelected>>", lambda e: self.refresh())
            self.filter_vars[field] = var
            col += 2

        date_field = _first_date_field(self.table, rows[0] if rows else {})
        self.year_field = date_field
        if date_field:
            years = sorted(
                {_field_year(r.get(date_field)) for r in rows if _field_year(r.get(date_field))},
                reverse=True,
            )
            ttk.Label(self.filter_frame, text="Year:").grid(
                row=0, column=col, sticky="w", padx=(0, 3)
            )
            combo = ttk.Combobox(
                self.filter_frame,
                textvariable=self.year_var,
                values=["All"] + years,
                state="readonly",
                width=9,
            )
            combo.grid(row=0, column=col + 1, sticky="w", padx=(0, 8))
            combo.bind("<<ComboboxSelected>>", lambda e: self.refresh())
            col += 2

        ttk.Button(
            self.filter_frame, text="Clear Filters", command=self.clear_filters
        ).grid(row=0, column=col, sticky="w", padx=3)

    def clear_filters(self):
        for var in self.filter_vars.values():
            var.set("All")
        self.year_var.set("All")
        self.refresh()

    def _heading_text(self, column):
        label = column.replace("_", " ").title()
        if self.sort_column == column:
            return f"{label} {'▼' if self.sort_descending else '▲'}"
        return label

    def _update_headings(self):
        for column in self.config_data["display"]:
            self.tree.heading(column, text=self._heading_text(column))

    @staticmethod
    def _natural_text_key(value):
        text = "" if value is None else str(value).strip().casefold()
        parts = re.split(r"(\d+)", text)
        return tuple(int(p) if p.isdigit() else p for p in parts)

    def _sort_key(self, row, column):
        app_module = sys.modules.get(self.winfo_toplevel().__class__.__module__)
        value = row.get(column)
        if column in getattr(app_module, "DATE_FIELDS", set()) or column in getattr(app_module, "YEAR_FIELDS", set()):
            return date_sort_key(value, present_is_latest=column in {"end_date", "end_year", "last_used"})
        if column in FLOAT_FIELDS | INT_FIELDS:
            return _to_float(value)
        if column in BOOL_FIELDS:
            return 1 if _truthy(value) else 0
        return self._natural_text_key(value)

    def sort_by(self, column):
        if column == self.sort_column:
            self.sort_descending = not self.sort_descending
        else:
            self.sort_column = column
            self.sort_descending = False
        self._update_headings()
        self.refresh()

    def _passes_filters(self, row):
        for field, var in self.filter_vars.items():
            selected = var.get()
            if selected != "All":
                actual = str(row.get(field) if row.get(field) is not None else "").strip()
                if actual != selected:
                    return False
        if getattr(self, "year_field", None) and self.year_var.get() != "All":
            if _field_year(row.get(self.year_field)) != self.year_var.get():
                return False
        return True

    def refresh(self):
        term = self.search_var.get().lower().strip()
        selected = set(self.tree.selection()) if hasattr(self, "tree") else set()
        self.tree.delete(*self.tree.get_children())

        all_rows = _rows(self.db, self.table)
        rows = []
        for row in all_rows:
            hay = " ".join(str(v or "") for v in row.values()).lower()
            if term and term not in hay:
                continue
            if not self._passes_filters(row):
                continue
            rows.append(row)

        if self.sort_column:
            populated = [r for r in rows if str(r.get(self.sort_column) or "").strip()]
            blanks = [r for r in rows if not str(r.get(self.sort_column) or "").strip()]
            populated.sort(
                key=lambda r: self._sort_key(r, self.sort_column),
                reverse=self.sort_descending,
            )
            rows = populated + blanks

        for row in rows:
            values = []
            for col in self.config_data["display"]:
                value = row.get(col, "")
                if col in BOOL_FIELDS:
                    value = "Yes" if _truthy(value) else "No"
                values.append("" if value is None else value)
            iid = str(row["id"])
            self.tree.insert("", "end", iid=iid, values=values)
            if iid in selected:
                self.tree.selection_add(iid)

        text = f"Showing {len(rows):,} of {len(all_rows):,} record(s)"
        self.count_label.configure(text=text)
        self.status_callback(f"{self.config_data['label']}: {text}")

    def selected_id(self):
        sel = self.tree.selection()
        return int(sel[0]) if sel else None

    def add(self):
        dlg = EnhancedRecordDialog(self, self.table, self.config_data)
        self.wait_window(dlg)
        if dlg.result is not None:
            self.db.insert_row(self.table, dlg.result)
            self._build_filters()
            self.refresh()
            self._refresh_root_dashboard()

    def edit(self):
        row_id = self.selected_id()
        if not row_id:
            messagebox.showinfo("Select Record", "Select a record to edit.", parent=self)
            return
        initial = self.db.get_row(self.table, row_id)
        dlg = EnhancedRecordDialog(self, self.table, self.config_data, initial)
        self.wait_window(dlg)
        if dlg.result is not None:
            self.db.update_row(self.table, row_id, dlg.result)
            self._build_filters()
            self.refresh()
            self._refresh_root_dashboard()

    def delete(self):
        row_id = self.selected_id()
        if not row_id:
            messagebox.showinfo("Select Record", "Select a record to delete.", parent=self)
            return
        if messagebox.askyesno("Delete Record", "Delete the selected record?", parent=self):
            self.db.delete_row(self.table, row_id)
            self._build_filters()
            self.refresh()
            self._refresh_root_dashboard()

    def _refresh_root_dashboard(self):
        root = _root_app(self)
        if hasattr(root, "refresh_dashboard"):
            try:
                root.refresh_dashboard()
            except Exception:
                pass

    def view_linked_records(self):
        row_id = self.selected_id()
        if not row_id:
            messagebox.showinfo(
                "Select Case", "Select a Case Index record first.", parent=self
            )
            return
        row = self.db.get_row("case_index", row_id)
        if row:
            CaseRelationshipWindow(_root_app(self), row)

    def normalize_casework_tools(self):
        rows = _rows(self.db, "casework")
        changed = 0
        for row in rows:
            normalized = "; ".join(normalize_tools(self.db, row.get("tools_used")))
            current = str(row.get("tools_used") or "").strip()
            if normalized and normalized != current:
                self.db.update_row("casework", int(row["id"]), {"tools_used": normalized})
                changed += 1
        messagebox.showinfo(
            "Tool Normalization",
            f"Normalized tool names in {changed} Case Work record(s).",
            parent=self,
        )
        self.refresh()


# ---------------------------------------------------------------------------
# Navigation / Global Search / Case Relationship Explorer
# ---------------------------------------------------------------------------

def navigate_to_record(app, table: str, row_id: int | None = None):
    tab = getattr(app, "record_tabs", {}).get(table)
    if tab is None:
        messagebox.showinfo(
            "Navigation",
            f"The {table.replace('_', ' ').title()} tab is not currently available.",
            parent=app,
        )
        return

    widget = tab
    while widget is not None and widget is not app:
        parent = getattr(widget, "master", None)
        if isinstance(parent, ttk.Notebook):
            try:
                parent.select(widget)
            except tk.TclError:
                pass
        widget = parent

    try:
        tab.refresh()
        if row_id is not None:
            iid = str(row_id)
            if tab.tree.exists(iid):
                tab.tree.selection_set(iid)
                tab.tree.focus(iid)
                tab.tree.see(iid)
    except Exception:
        pass


def _row_date(row: dict[str, Any]) -> str:
    field = _first_date_field("", row)
    return str(row.get(field) or "") if field else ""


def _row_summary(app_module, table: str, row: dict[str, Any]) -> str:
    config = getattr(app_module, "TABLE_CONFIG", {}).get(table, {})
    fields = config.get("display", [])
    pieces = []
    for field in fields:
        value = str(row.get(field) or "").strip()
        if value and value not in pieces:
            pieces.append(value)
        if len(" — ".join(pieces)) > 140:
            break
    if not pieces:
        for field, value in row.items():
            if field in {"id", "profile_id", "sort_order", "case_id"}:
                continue
            text = str(value or "").strip()
            if text:
                pieces.append(text)
            if len(pieces) >= 3:
                break
    return " — ".join(pieces)[:220]


class GlobalSearchDialog(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.app = parent
        self.db = parent.db
        self.title("Global Professional Record Search")
        self.geometry("1050x650")
        self.minsize(800, 500)
        self.transient(parent)

        top = ttk.Frame(self, padding=10)
        top.pack(fill="x")
        ttk.Label(top, text="Search all professional records:", style="Title.TLabel").pack(anchor="w")
        row = ttk.Frame(top)
        row.pack(fill="x", pady=(8, 0))
        self.query = tk.StringVar()
        entry = ttk.Entry(row, textvariable=self.query)
        entry.pack(side="left", fill="x", expand=True)
        ttk.Button(row, text="Search", command=self.search).pack(side="left", padx=(6, 0))
        ttk.Button(row, text="Open Selected", command=self.open_selected).pack(side="left", padx=(6, 0))

        cols = ("section", "date", "summary")
        self.tree = ttk.Treeview(self, columns=cols, show="headings", selectmode="browse")
        self.tree.heading("section", text="Section")
        self.tree.heading("date", text="Date")
        self.tree.heading("summary", text="Summary")
        self.tree.column("section", width=190, stretch=False)
        self.tree.column("date", width=110, stretch=False)
        self.tree.column("summary", width=700, stretch=True)
        scroll = ttk.Scrollbar(self, orient="vertical", command=self.tree.yview)
        self.tree.configure(yscrollcommand=scroll.set)
        self.tree.pack(side="left", fill="both", expand=True, padx=(10, 0), pady=(0, 10))
        scroll.pack(side="right", fill="y", padx=(0, 10), pady=(0, 10))
        self.tree.bind("<Double-1>", lambda e: self.open_selected())
        self.results = {}
        self.status = ttk.Label(top, text="Enter at least 2 characters.")
        self.status.pack(anchor="w", pady=(5, 0))
        entry.bind("<Return>", lambda e: self.search())
        entry.focus_set()

    def search(self):
        q = self.query.get().strip().casefold()
        self.tree.delete(*self.tree.get_children())
        self.results.clear()
        if len(q) < 2:
            self.status.configure(text="Enter at least 2 characters.")
            return

        app_module = sys.modules.get(self.app.__class__.__module__)
        count = 0
        limit = 500
        for table in database_module.DATA_TABLES:
            if table not in database_module.TABLE_FIELDS:
                continue
            for row in _rows(self.db, table):
                hay = " ".join(str(v or "") for v in row.values()).casefold()
                if q not in hay:
                    continue
                label = getattr(app_module, "TABLE_CONFIG", {}).get(table, {}).get(
                    "label", table.replace("_", " ").title()
                )
                iid = f"r{count}"
                self.tree.insert(
                    "", "end", iid=iid,
                    values=(label, _row_date(row), _row_summary(app_module, table, row)),
                )
                self.results[iid] = (table, int(row["id"]))
                count += 1
                if count >= limit:
                    break
            if count >= limit:
                break
        suffix = " (result limit reached)" if count >= limit else ""
        self.status.configure(text=f"{count:,} result(s){suffix}")

    def open_selected(self):
        sel = self.tree.selection()
        if not sel:
            return
        target = self.results.get(sel[0])
        if target:
            navigate_to_record(self.app, target[0], target[1])
            self.destroy()


class CaseRelationshipWindow(tk.Toplevel):
    def __init__(self, parent, case_row: dict[str, Any]):
        super().__init__(parent)
        self.app = parent
        self.db = parent.db
        self.case_row = case_row
        self.case_number = _case_number(case_row.get("case_number"))
        self.title(f"Linked Records — {self.case_number}")
        self.geometry("1050x620")
        self.minsize(800, 480)
        self.transient(parent)

        header = ttk.Frame(self, padding=10)
        header.pack(fill="x")
        ttk.Label(
            header,
            text=f"Case / Reference: {self.case_number}",
            style="Title.TLabel",
        ).pack(anchor="w")
        metadata = " | ".join(
            x for x in [
                str(case_row.get("agency") or "").strip(),
                str(case_row.get("case_type") or "").strip(),
                str(case_row.get("status") or "").strip(),
            ] if x
        )
        if metadata:
            ttk.Label(header, text=metadata).pack(anchor="w", pady=(3, 0))

        cols = ("section", "date", "summary")
        self.tree = ttk.Treeview(self, columns=cols, show="headings", selectmode="browse")
        for col, title, width in [
            ("section", "Section", 200),
            ("date", "Date", 110),
            ("summary", "Linked Record", 700),
        ]:
            self.tree.heading(col, text=title)
            self.tree.column(col, width=width, stretch=(col == "summary"))
        self.tree.pack(fill="both", expand=True, padx=10, pady=(0, 10))
        self.targets = {}
        self.tree.bind("<Double-1>", lambda e: self.open_selected())
        ttk.Button(
            header, text="Open Selected Record", command=self.open_selected
        ).pack(anchor="e")

        self.populate()

    def populate(self):
        app_module = sys.modules.get(self.app.__class__.__module__)
        count = 0
        for table in LINKED_TABLES:
            if table not in database_module.TABLE_FIELDS:
                continue
            for row in _rows(self.db, table):
                if _case_number(row.get("case_number")).casefold() != self.case_number.casefold():
                    continue
                label = getattr(app_module, "TABLE_CONFIG", {}).get(table, {}).get(
                    "label", table.replace("_", " ").title()
                )
                iid = f"r{count}"
                self.tree.insert(
                    "", "end", iid=iid,
                    values=(label, _row_date(row), _row_summary(app_module, table, row)),
                )
                self.targets[iid] = (table, int(row["id"]))
                count += 1

    def open_selected(self):
        sel = self.tree.selection()
        if not sel:
            return
        target = self.targets.get(sel[0])
        if target:
            navigate_to_record(self.app, target[0], target[1])


# ---------------------------------------------------------------------------
# Annual statistics
# ---------------------------------------------------------------------------

YEAR_TABLE_FIELDS = {
    "casework": "examination_date",
    "forensic_reports": "report_date",
    "peer_reviews": "review_date",
    "testimony": "testimony_date",
    "court_qualifications": "qualification_date",
    "training": "attended_date",
    "certifications": "earned_date",
    "teaching": "start_date",
    "presentations": "presentation_date",
    "publications": "publication_date",
    "validations": "validation_date",
    "procedures": "effective_date",
    "mentoring": "start_date",
    "projects": "start_date",
    "achievements": "achievement_date",
}


def annual_rows(db, table: str, year: int) -> list[dict[str, Any]]:
    field = YEAR_TABLE_FIELDS.get(table)
    if not field:
        return []
    return [
        row for row in _rows(db, table)
        if _field_year(row.get(field)) == str(year)
    ]


def annual_activity(db, year: int) -> dict[str, Any]:
    casework = annual_rows(db, "casework", year)
    reports = annual_rows(db, "forensic_reports", year)
    reviews = annual_rows(db, "peer_reviews", year)
    testimony = annual_rows(db, "testimony", year)
    qualifications = annual_rows(db, "court_qualifications", year)
    training = annual_rows(db, "training", year)
    teaching = annual_rows(db, "teaching", year)
    presentations = annual_rows(db, "presentations", year)
    publications = annual_rows(db, "publications", year)
    validations = annual_rows(db, "validations", year)
    procedures = annual_rows(db, "procedures", year)
    mentoring = annual_rows(db, "mentoring", year)
    projects = annual_rows(db, "projects", year)

    unique_cases = {
        _case_number(r.get("case_number")).casefold()
        for r in casework if _case_number(r.get("case_number"))
    }
    tools = canonical_tool_counter(db, casework)
    device_types = Counter(
        str(r.get("device_type") or "").strip()
        for r in casework if str(r.get("device_type") or "").strip()
    )
    acquisitions = Counter(
        str(r.get("acquisition_type") or "").strip()
        for r in casework if str(r.get("acquisition_type") or "").strip()
    )

    return {
        "year": year,
        "casework": len(casework),
        "unique_cases": len(unique_cases),
        "exam_hours": sum(_to_float(r.get("hours")) for r in casework),
        "reports": len(reports),
        "reports_authored": sum(
            1 for r in reports if "author" in str(r.get("role") or "").casefold()
        ),
        "reviews": len(reviews),
        "review_hours": sum(_to_float(r.get("hours")) for r in reviews),
        "testimony": len(testimony),
        "expert_testimony": sum(
            1 for r in testimony if str(r.get("witness_type") or "") == "Expert Witness"
        ),
        "qualifications": len(qualifications),
        "training_records": len(training),
        "training_hours": sum(_to_float(r.get("hours")) for r in training),
        "cpe": sum(_to_float(r.get("cpe_credits")) for r in training),
        "teaching_records": len(teaching),
        "teaching_hours": sum(_to_float(r.get("hours")) for r in teaching),
        "students": sum(_to_int(r.get("students")) for r in teaching),
        "presentations": len(presentations),
        "publications": len(publications),
        "validations": len(validations),
        "procedures": len(procedures),
        "mentoring": len(mentoring),
        "projects": len(projects),
        "device_types": device_types,
        "acquisitions": acquisitions,
        "tools": tools,
    }


# ---------------------------------------------------------------------------
# Dashboard drill-down + alerts
# ---------------------------------------------------------------------------

def _expiration_info(raw: Any):
    text = str(raw or "").strip()
    if not text or text == "Present":
        return None
    try:
        norm = normalize_date(text, allow_year_only=True)
        parts = [int(p) for p in norm.split("-")]
        y = parts[0]
        m = parts[1] if len(parts) > 1 else 12
        d = parts[2] if len(parts) > 2 else 28
        d = min(d, 28)
        days = (date(y, m, d) - date.today()).days
        if days < 0:
            return ("red", f"EXPIRED {-days} day(s) ago")
        if days <= 90:
            return ("red", f"expires in {days} day(s)")
        if days <= 180:
            return ("orange", f"expires in {days} day(s)")
        if days <= 365:
            return ("yellow", f"expires in {days} day(s)")
        return ("normal", f"expires in {days} day(s)")
    except Exception:
        return ("normal", f"expiration: {text}")


def certification_alerts(db) -> list[tuple[str, str]]:
    alerts = []
    for row in _rows(db, "certifications"):
        name = str(row.get("certification") or "Certification")
        expiration = _expiration_info(row.get("expiration_date"))
        req = _to_float(row.get("cpe_required"))
        earned = _to_float(row.get("cpe_earned"))
        if expiration and expiration[0] in {"red", "orange", "yellow"}:
            alerts.append((expiration[0], f"{name}: {expiration[1]}"))
        if req > 0 and earned < req:
            remaining = req - earned
            severity = "orange" if earned / req < 0.5 else "yellow"
            alerts.append(
                (severity, f"{name}: CPE/CE {earned:g}/{req:g} — {remaining:g} remaining")
            )
    return alerts


def enhanced_refresh_dashboard(self):
    metrics = pt.professional_metrics(self.db) if pt else {}
    case = metrics.get("casework", {})
    reports = metrics.get("reports", {})
    reviews = metrics.get("reviews", {})
    court = metrics.get("court", {})
    training = metrics.get("training", {})
    current = annual_activity(self.db, date.today().year)

    for w in self.metrics_frame.winfo_children():
        w.destroy()

    cards = [
        ("Cases", int(case.get("unique_cases", 0)), "case_index"),
        ("Examinations", int(case.get("examinations", 0)), "casework"),
        ("Reports", int(reports.get("total", 0)), "forensic_reports"),
        ("Peer Reviews", int(reviews.get("total", 0)), "peer_reviews"),
        ("Testimony", int(court.get("testimony_total", 0)), "testimony"),
        ("Training Hrs", float(training.get("hours", 0)), "training"),
        (f"{date.today().year} Exams", current["casework"], "casework"),
        (f"{date.today().year} Reports", current["reports"], "forensic_reports"),
    ]

    for i, (label, value, table) in enumerate(cards):
        box = ttk.LabelFrame(self.metrics_frame, text=label, padding=7)
        box.grid(row=i // 4, column=i % 4, padx=4, pady=4, sticky="nsew")
        display = f"{value:,.1f}" if isinstance(value, float) else f"{value:,}"
        btn = ttk.Button(
            box,
            text=display,
            command=lambda t=table: navigate_to_record(self, t),
        )
        btn.pack(fill="both", expand=True)
        self.metrics_frame.columnconfigure(i % 4, weight=1)

    # Summary and color-coded alerts.
    self.summary_text.config(state="normal")
    self.summary_text.delete("1.0", "end")
    self.summary_text.tag_configure("heading", font=("Segoe UI", 10, "bold"))
    self.summary_text.tag_configure("red", foreground="#c62828")
    self.summary_text.tag_configure("orange", foreground="#e65100")
    self.summary_text.tag_configure("yellow", foreground="#9a7600")

    self.summary_text.insert("end", "Current-year activity\n", "heading")
    lines = [
        f"{current['year']} unique cases: {current['unique_cases']:,}",
        f"{current['year']} examinations: {current['casework']:,}",
        f"{current['year']} examination hours: {current['exam_hours']:,.1f}",
        f"{current['year']} reports: {current['reports']:,}",
        f"{current['year']} peer reviews: {current['reviews']:,}",
        f"{current['year']} testimony appearances: {current['testimony']:,}",
        f"{current['year']} training hours: {current['training_hours']:,.1f}",
        f"{current['year']} instruction hours: {current['teaching_hours']:,.1f}",
    ]
    self.summary_text.insert("end", "\n".join(lines) + "\n\n")

    self.summary_text.insert("end", "Certification / CPE alerts\n", "heading")
    alerts = certification_alerts(self.db)
    if alerts:
        for severity, text in alerts:
            self.summary_text.insert("end", f"• {text}\n", severity)
    else:
        self.summary_text.insert("end", "No urgent certification/CPE alerts.\n")

    self.summary_text.insert("end", "\nData protection\n", "heading")
    self.summary_text.insert(
        "end",
        f"Schema version: {get_schema_version(self.db)}\n"
        f"Automatic backups: daily, newest {BACKUP_RETENTION} retained\n"
        f"Database: {self.db_path}\n",
    )
    self.summary_text.config(state="disabled")

    # Keep a simple record-count chart.
    counts = [
        ("Case Work", int(case.get("examinations", 0))),
        ("Reports", int(reports.get("total", 0))),
        ("Reviews", int(reviews.get("total", 0))),
        ("Testimony", int(court.get("testimony_total", 0))),
        ("Training", int(training.get("records", 0))),
        ("Certs", int(training.get("certifications", 0))),
    ]
    self.dashboard_chart.delete("all")
    self.dashboard_chart.update_idletasks()
    width = max(self.dashboard_chart.winfo_width(), 360)
    height = max(self.dashboard_chart.winfo_height(), 260)
    max_value = max([v for _, v in counts] + [1])
    left, top, bottom = 78, 12, height - 22
    usable = max(120, width - left - 35)
    bar_h = max(12, min(26, int((bottom - top) / len(counts) - 5)))
    chart_fg = "#f2f3f5" if getattr(self, "theme_name", "light") == "dark" else "#212529"
    for i, (label, value) in enumerate(counts):
        y = top + i * ((bottom - top) / len(counts))
        x2 = left + usable * (value / max_value)
        self.dashboard_chart.create_text(
            left - 7, y + bar_h / 2, text=label, anchor="e", fill=chart_fg
        )
        self.dashboard_chart.create_rectangle(
            left, y, x2, y + bar_h, fill="#1f4e79", outline=""
        )
        self.dashboard_chart.create_text(
            x2 + 5, y + bar_h / 2, text=str(value), anchor="w", fill=chart_fg
        )


# ---------------------------------------------------------------------------
# Export presets
# ---------------------------------------------------------------------------

BUILTIN_PRESETS = {
    "Standard CV": {
        "summary": 1, "core_training": 1, "employment": 1, "teaching": 1,
        "organizations": 1, "certifications": 1, "skills": 1, "education": 1,
        "testimony": 1, "achievements": 1,
        "forensic_case_summary": 1, "forensic_report_summary": 1,
        "court_summary": 1, "training_summary": 1, "tool_summary": 1,
    },
    "Expert Witness CV": {
        "summary": 1, "employment": 1, "education": 1, "certifications": 1,
        "skills": 1, "testimony": 1, "core_training": 1, "teaching": 1,
        "forensic_case_summary": 1, "forensic_report_summary": 1,
        "court_summary": 1, "training_summary": 1, "tool_summary": 1,
        "quality_summary": 1,
    },
    "Voir Dire Package": {
        "summary": 1, "employment": 1, "education": 1, "certifications": 1,
        "skills": 1, "testimony": 1, "core_training": 1,
        "forensic_case_summary": 1, "forensic_report_summary": 1,
        "court_summary": 1, "training_summary": 1, "tool_summary": 1,
        "quality_summary": 1, "detail_court_qualifications": 1,
        "detail_testimony": 1, "detail_tool_experience": 1,
        "detail_certifications": 1,
    },
    "Full Professional Record": {
        "__all_cv": 1, "__all_summary": 1, "__all_detail": 1,
    },
    "Annual Review": {
        "summary": 1, "employment": 1, "achievements": 1,
        "forensic_case_summary": 1, "forensic_report_summary": 1,
        "training_summary": 1, "teaching_summary": 1,
        "quality_summary": 1, "publication_summary": 1,
        "leadership_summary": 1,
    },
    "Training / CPE Report": {
        "certifications": 1, "core_training": 1,
        "training_summary": 1, "detail_training": 1,
        "detail_certifications": 1,
    },
}


def _preset_path(app_module) -> Path:
    return app_module.portable_data_dir() / "export_presets.json"


def load_custom_presets(app_module) -> dict[str, dict[str, int]]:
    try:
        data = json.loads(_preset_path(app_module).read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def save_custom_presets(app_module, data):
    _preset_path(app_module).write_text(
        json.dumps(data, indent=2, sort_keys=True),
        encoding="utf-8",
    )


def apply_preset(app, name: str):
    app_module = sys.modules.get(app.__class__.__module__)
    custom = load_custom_presets(app_module)
    preset = BUILTIN_PRESETS.get(name) or custom.get(name)
    if not preset:
        return

    for var in app.cv_options.values():
        var.set(0)

    if preset.get("__all_cv") and pt:
        for key, _, _ in pt.CV_OPTIONS:
            if key in app.cv_options:
                app.cv_options[key].set(1)
    if preset.get("__all_summary") and pt:
        for key, _, _ in pt.SUMMARY_OPTIONS:
            if key in app.cv_options:
                app.cv_options[key].set(1)
    if preset.get("__all_detail") and pt:
        for key, _, _ in pt.DETAIL_OPTIONS:
            if key in app.cv_options:
                app.cv_options[key].set(1)

    for key, value in preset.items():
        if key.startswith("__"):
            continue
        if key in app.cv_options:
            app.cv_options[key].set(int(bool(value)))

    # Privacy fields remain OFF unless a custom preset explicitly turns them on.
    if pt:
        for key, _, _ in pt.PRIVACY_OPTIONS:
            if key in app.cv_options and key not in preset:
                app.cv_options[key].set(0)

    app.set_status(f"Export preset applied: {name}")


def save_current_preset(app):
    app_module = sys.modules.get(app.__class__.__module__)
    name = simpledialog.askstring(
        "Save Export Preset",
        "Preset name:",
        parent=app,
    )
    if not name:
        return
    name = name.strip()
    if name in BUILTIN_PRESETS:
        messagebox.showerror(
            "Preset Name",
            "That name is reserved for a built-in preset.",
            parent=app,
        )
        return
    custom = load_custom_presets(app_module)
    custom[name] = {key: int(bool(var.get())) for key, var in app.cv_options.items()}
    save_custom_presets(app_module, custom)
    refresh_preset_combo(app)
    app.preset_var.set(name)
    app.set_status(f"Saved export preset: {name}")


def delete_custom_preset(app):
    app_module = sys.modules.get(app.__class__.__module__)
    name = getattr(app, "preset_var", tk.StringVar()).get()
    custom = load_custom_presets(app_module)
    if name not in custom:
        messagebox.showinfo(
            "Delete Preset",
            "Select a custom preset. Built-in presets cannot be deleted.",
            parent=app,
        )
        return
    if messagebox.askyesno("Delete Preset", f'Delete preset "{name}"?', parent=app):
        custom.pop(name, None)
        save_custom_presets(app_module, custom)
        refresh_preset_combo(app)
        app.preset_var.set("Standard CV")


def refresh_preset_combo(app):
    app_module = sys.modules.get(app.__class__.__module__)
    names = list(BUILTIN_PRESETS) + sorted(load_custom_presets(app_module), key=str.casefold)
    app.preset_combo.configure(values=names)


def add_preset_toolbar(app, generate_tab):
    app_module = sys.modules.get(app.__class__.__module__)
    bar = ttk.LabelFrame(generate_tab, text="Output Presets & Specialty Reports", padding=8)
    children = generate_tab.winfo_children()
    if children:
        bar.pack(fill="x", padx=16, pady=(8, 2), before=children[0])
    else:
        bar.pack(fill="x", padx=16, pady=(8, 2))

    ttk.Label(bar, text="Preset:").pack(side="left")
    app.preset_var = tk.StringVar(value="Standard CV")
    app.preset_combo = ttk.Combobox(
        bar, textvariable=app.preset_var, state="readonly", width=27
    )
    app.preset_combo.pack(side="left", padx=(5, 4))
    refresh_preset_combo(app)
    ttk.Button(
        bar, text="Apply", command=lambda: apply_preset(app, app.preset_var.get())
    ).pack(side="left", padx=2)
    ttk.Button(
        bar, text="Save Current...", command=lambda: save_current_preset(app)
    ).pack(side="left", padx=2)
    ttk.Button(
        bar, text="Delete Custom", command=lambda: delete_custom_preset(app)
    ).pack(side="left", padx=2)

    ttk.Separator(bar, orient="vertical").pack(side="left", fill="y", padx=8)
    ttk.Button(
        bar, text="Expert Qualification Report...",
        command=lambda: generate_expert_report_interactive(app),
    ).pack(side="left", padx=2)
    ttk.Button(
        bar, text="Annual Activity Report...",
        command=lambda: generate_annual_report_interactive(app),
    ).pack(side="left", padx=2)

    apply_preset(app, "Standard CV")


# ---------------------------------------------------------------------------
# Specialty Word + PDF reports
# ---------------------------------------------------------------------------

def _profile_title(db) -> str:
    p = db.get_profile()
    return (
        p.get("preferred_name")
        or p.get("full_name")
        or p.get("profile_name")
        or "Examiner"
    )


def _report_doc_header(doc, title: str, subtitle: str):
    from docx.shared import Pt
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(18)
    if subtitle:
        p2 = doc.add_paragraph()
        r2 = p2.add_run(subtitle)
        r2.italic = True
        r2.font.size = Pt(10)


def _doc_bullets(doc, lines):
    for line in lines:
        doc.add_paragraph(str(line), style="List Bullet")


def generate_expert_docx(db, path: Path) -> Path:
    from docx import Document

    m = pt.professional_metrics(db) if pt else {}
    case = m.get("casework", {})
    reports = m.get("reports", {})
    reviews = m.get("reviews", {})
    court = m.get("court", {})
    training = m.get("training", {})
    quality = m.get("quality", {})
    teaching = m.get("teaching", {})

    doc = Document()
    _report_doc_header(
        doc,
        "Expert Qualification / Voir Dire Summary",
        f"{_profile_title(db)} — generated {date.today().isoformat()}",
    )

    doc.add_heading("Professional Experience Metrics", level=1)
    _doc_bullets(doc, [
        f"Unique cases documented: {int(case.get('unique_cases', 0)):,}",
        f"Digital forensic examinations documented: {int(case.get('examinations', 0)):,}",
        f"Documented examination hours: {float(case.get('hours', 0)):,.1f}",
        f"Forensic reports tracked: {int(reports.get('total', 0)):,}",
        f"Peer / technical reviews tracked: {int(reviews.get('total', 0)):,}",
        f"Testimony appearances tracked: {int(court.get('testimony_total', 0)):,}",
        f"Expert-witness appearances: {int(court.get('expert', 0)):,}",
        f"Court qualification records: {int(court.get('qualification_records', 0)):,}",
        f"Training hours documented: {float(training.get('hours', 0)):,.1f}",
        f"Certifications tracked: {int(training.get('certifications', 0)):,}",
        f"Instruction hours documented: {float(teaching.get('hours', 0)):,.1f}",
        f"Tool validation/testing records: {int(quality.get('validations', 0)):,}",
    ])

    casework = _rows(db, "casework")
    tools = canonical_tool_counter(db, casework)
    if tools:
        doc.add_heading("Primary Forensic Tools Used in Documented Case Work", level=1)
        _doc_bullets(doc, [f"{name}: {count:,} examination(s)" for name, count in tools.most_common(15)])

    quals = _rows(db, "court_qualifications")
    if quals:
        doc.add_heading("Court Qualification History", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        for i, text in enumerate(["Date", "Court", "Jurisdiction", "Area", "Result"]):
            table.rows[0].cells[i].text = text
        for row in quals:
            vals = [
                row.get("qualification_date", ""),
                row.get("court", ""),
                row.get("jurisdiction", ""),
                row.get("qualification_area", ""),
                row.get("ruling", ""),
            ]
            cells = table.add_row().cells
            for i, value in enumerate(vals):
                cells[i].text = str(value or "")

    certs = _rows(db, "certifications")
    if certs:
        doc.add_heading("Certifications", level=1)
        for row in certs:
            name = row.get("certification") or ""
            issuer = row.get("issuing_organization") or ""
            status = row.get("status") or ""
            doc.add_paragraph(f"{name} — {issuer} — {status}", style="List Bullet")

    doc.add_paragraph(
        "This report summarizes records entered into Forensic CV Manager Professional. "
        "It intentionally omits case numbers and evidence identifiers."
    )
    doc.save(path)
    return path


def generate_annual_docx(db, path: Path, year: int) -> Path:
    from docx import Document
    a = annual_activity(db, year)
    doc = Document()
    _report_doc_header(
        doc,
        f"{year} Professional Activity Report",
        f"{_profile_title(db)} — generated {date.today().isoformat()}",
    )
    doc.add_heading("Activity Summary", level=1)
    _doc_bullets(doc, [
        f"Unique cases: {a['unique_cases']:,}",
        f"Digital forensic examinations: {a['casework']:,}",
        f"Examination hours: {a['exam_hours']:,.1f}",
        f"Forensic reports: {a['reports']:,}",
        f"Reports authored/co-authored: {a['reports_authored']:,}",
        f"Peer / technical reviews: {a['reviews']:,}",
        f"Peer review hours: {a['review_hours']:,.1f}",
        f"Testimony appearances: {a['testimony']:,}",
        f"Expert-witness appearances: {a['expert_testimony']:,}",
        f"Court qualification events: {a['qualifications']:,}",
        f"Training hours: {a['training_hours']:,.1f}",
        f"CPE / CE credits: {a['cpe']:,.1f}",
        f"Instruction hours: {a['teaching_hours']:,.1f}",
        f"Students / attendees documented: {a['students']:,}",
        f"Presentations: {a['presentations']:,}",
        f"Publications / research: {a['publications']:,}",
        f"Tool validation/testing records: {a['validations']:,}",
        f"SOP / policy records: {a['procedures']:,}",
        f"Mentoring / supervision records: {a['mentoring']:,}",
        f"Major projects: {a['projects']:,}",
    ])
    for heading, counter in [
        ("Device Types", a["device_types"]),
        ("Acquisition Methods", a["acquisitions"]),
        ("Forensic Tools Used", a["tools"]),
    ]:
        if counter:
            doc.add_heading(heading, level=1)
            _doc_bullets(doc, [f"{k}: {v:,}" for k, v in counter.most_common()])
    doc.add_paragraph(
        "Counts are based on records with dates in the selected calendar year. "
        "Case numbers and evidence identifiers are not included."
    )
    doc.save(path)
    return path


def _pdf_lines(path: Path, title: str, subtitle: str, sections: list[tuple[str, list[str]]]):
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib import colors

    styles = getSampleStyleSheet()
    h = ParagraphStyle(
        "V25H", parent=styles["Heading1"], fontSize=12, leading=14,
        textColor=colors.HexColor("#1F4E79"), spaceBefore=8, spaceAfter=4,
    )
    body = ParagraphStyle(
        "V25Body", parent=styles["BodyText"], fontSize=9, leading=11, spaceAfter=2,
    )
    doc = SimpleDocTemplate(
        str(path), pagesize=LETTER,
        leftMargin=.65 * inch, rightMargin=.65 * inch,
        topMargin=.6 * inch, bottomMargin=.6 * inch,
    )
    story = [
        Paragraph(html.escape(title), styles["Title"]),
        Paragraph(html.escape(subtitle), body),
        Spacer(1, 8),
    ]
    for heading, lines in sections:
        story.append(Paragraph(html.escape(heading), h))
        for line in lines:
            story.append(Paragraph("• " + html.escape(str(line)), body))
    doc.build(story)


def generate_expert_pdf(db, path: Path):
    m = pt.professional_metrics(db) if pt else {}
    case = m.get("casework", {})
    court = m.get("court", {})
    reports = m.get("reports", {})
    reviews = m.get("reviews", {})
    training = m.get("training", {})
    tools = canonical_tool_counter(db, _rows(db, "casework"))

    sections = [
        ("Professional Experience Metrics", [
            f"Unique cases documented: {int(case.get('unique_cases', 0)):,}",
            f"Forensic examinations documented: {int(case.get('examinations', 0)):,}",
            f"Examination hours: {float(case.get('hours', 0)):,.1f}",
            f"Forensic reports tracked: {int(reports.get('total', 0)):,}",
            f"Peer / technical reviews: {int(reviews.get('total', 0)):,}",
            f"Expert-witness appearances: {int(court.get('expert', 0)):,}",
            f"Court qualification records: {int(court.get('qualification_records', 0)):,}",
            f"Training hours: {float(training.get('hours', 0)):,.1f}",
            f"Certifications: {int(training.get('certifications', 0)):,}",
        ]),
        ("Primary Forensic Tools", [
            f"{name}: {count:,} examination(s)"
            for name, count in tools.most_common(15)
        ]),
    ]
    _pdf_lines(
        path,
        "Expert Qualification / Voir Dire Summary",
        f"{_profile_title(db)} — generated {date.today().isoformat()}",
        sections,
    )
    return path


def generate_annual_pdf(db, path: Path, year: int):
    a = annual_activity(db, year)
    sections = [
        ("Activity Summary", [
            f"Unique cases: {a['unique_cases']:,}",
            f"Forensic examinations: {a['casework']:,}",
            f"Examination hours: {a['exam_hours']:,.1f}",
            f"Forensic reports: {a['reports']:,}",
            f"Peer / technical reviews: {a['reviews']:,}",
            f"Testimony appearances: {a['testimony']:,}",
            f"Expert-witness appearances: {a['expert_testimony']:,}",
            f"Court qualification events: {a['qualifications']:,}",
            f"Training hours: {a['training_hours']:,.1f}",
            f"CPE / CE credits: {a['cpe']:,.1f}",
            f"Instruction hours: {a['teaching_hours']:,.1f}",
            f"Presentations: {a['presentations']:,}",
            f"Publications / research: {a['publications']:,}",
            f"Tool validations: {a['validations']:,}",
            f"SOP / policy records: {a['procedures']:,}",
            f"Mentoring records: {a['mentoring']:,}",
            f"Major projects: {a['projects']:,}",
        ]),
        ("Device Types", [f"{k}: {v:,}" for k, v in a["device_types"].most_common()]),
        ("Acquisition Methods", [f"{k}: {v:,}" for k, v in a["acquisitions"].most_common()]),
        ("Forensic Tools Used", [f"{k}: {v:,}" for k, v in a["tools"].most_common()]),
    ]
    _pdf_lines(
        path,
        f"{year} Professional Activity Report",
        f"{_profile_title(db)} — generated {date.today().isoformat()}",
        sections,
    )
    return path


def generate_expert_report_interactive(app):
    app_module = sys.modules.get(app.__class__.__module__)
    base = _profile_title(app.db).replace(" ", "_") + "_Expert_Qualification_Report.docx"
    out = filedialog.asksaveasfilename(
        title="Save Expert Qualification Report",
        defaultextension=".docx",
        initialdir=str(app_module.portable_resume_dir()),
        initialfile=base,
        filetypes=[("Word Document", "*.docx")],
    )
    if not out:
        return
    try:
        docx = generate_expert_docx(app.db, Path(out))
        pdf = generate_expert_pdf(app.db, Path(out).with_suffix(".pdf"))
        messagebox.showinfo(
            "Expert Report",
            f"Created:\n{docx.name}\n{pdf.name}\n\n"
            "Case numbers and evidence identifiers were omitted.",
            parent=app,
        )
        app.set_status(f"Expert qualification report generated: {docx}")
    except Exception as exc:
        messagebox.showerror("Expert Report", str(exc), parent=app)


def generate_annual_report_interactive(app):
    app_module = sys.modules.get(app.__class__.__module__)
    year = simpledialog.askinteger(
        "Annual Activity Report",
        "Calendar year:",
        initialvalue=date.today().year,
        minvalue=1900,
        maxvalue=2100,
        parent=app,
    )
    if not year:
        return
    base = _profile_title(app.db).replace(" ", "_") + f"_{year}_Professional_Activity.docx"
    out = filedialog.asksaveasfilename(
        title=f"Save {year} Professional Activity Report",
        defaultextension=".docx",
        initialdir=str(app_module.portable_resume_dir()),
        initialfile=base,
        filetypes=[("Word Document", "*.docx")],
    )
    if not out:
        return
    try:
        docx = generate_annual_docx(app.db, Path(out), year)
        pdf = generate_annual_pdf(app.db, Path(out).with_suffix(".pdf"), year)
        messagebox.showinfo(
            "Annual Activity Report",
            f"Created:\n{docx.name}\n{pdf.name}",
            parent=app,
        )
        app.set_status(f"{year} professional activity report generated")
    except Exception as exc:
        messagebox.showerror("Annual Activity Report", str(exc), parent=app)


# ---------------------------------------------------------------------------
# About / Changelog
# ---------------------------------------------------------------------------

def professional_about(self):
    upstream_version = ""
    try:
        import version
        upstream_version = getattr(version, "__version__", PRO_VERSION)
    except Exception:
        upstream_version = PRO_VERSION

    messagebox.showinfo(
        "About Forensic CV Manager Professional",
        f"Forensic CV Manager — {PRO_EDITION}\n"
        f"Version {PRO_VERSION}\n\n"
        "Expanded professional-history, digital-forensics experience, "
        "credential, testimony, quality-assurance, and reporting system.\n\n"
        f"Database schema: {get_schema_version(self.db)}\n"
        f"Automatic backups: daily / retain {BACKUP_RETENTION}\n"
        f"Application version metadata: {upstream_version}\n\n"
        "Based on Forensic CV Manager. Keep real case information only in "
        "storage environments authorized for that data.",
        parent=self,
    )


def show_changelog(self):
    win = tk.Toplevel(self)
    win.title(f"Professional Edition Changelog — {PRO_VERSION}")
    win.geometry("820x620")
    win.transient(self)
    frame = ttk.Frame(win, padding=10)
    frame.pack(fill="both", expand=True)
    ttk.Label(frame, text="Changelog", style="Title.TLabel").pack(anchor="w", pady=(0, 8))
    text = tk.Text(frame, wrap="word", padx=12, pady=12)
    scroll = ttk.Scrollbar(frame, orient="vertical", command=text.yview)
    text.configure(yscrollcommand=scroll.set)
    text.pack(side="left", fill="both", expand=True)
    scroll.pack(side="right", fill="y")
    text.insert("1.0", CHANGELOG)
    text.configure(state="disabled")


# ---------------------------------------------------------------------------
# App patch installation
# ---------------------------------------------------------------------------

_INSTALLED = False


def install_extensions(App):
    """Install the v2.5 upgrade after App and Professional Tracking are defined."""
    global _INSTALLED
    if _INSTALLED:
        return

    app_module = sys.modules.get(App.__module__)
    if app_module is None:
        raise RuntimeError("Could not resolve app module.")

    install_database_extensions()
    extend_table_config(app_module.TABLE_CONFIG, app_module.DATE_FIELDS, app_module.YEAR_FIELDS)
    install_auto_backup(app_module)

    # Global classes are resolved at runtime by the App grouping code.
    app_module.RecordDialog = EnhancedRecordDialog
    app_module.RecordsTab = EnhancedRecordsTab

    # Normalize tool aliases inside the existing Professional Tracking metrics too,
    # so normal CV/summary output does not split AXIOM/Magnet Axiom/etc.
    if pt and not getattr(pt, "_v25_metrics_wrapped", False):
        original_professional_metrics = pt.professional_metrics

        def normalized_professional_metrics(db):
            data = original_professional_metrics(db)
            counter = canonical_tool_counter(db, _rows(db, "casework"))
            try:
                data["casework"]["tools"] = counter
                data["tools"]["casework_tools"] = counter
            except Exception:
                pass
            return data

        pt.professional_metrics = normalized_professional_metrics
        pt._v25_metrics_wrapped = True

    # Dashboard replacement.
    App.refresh_dashboard = enhanced_refresh_dashboard

    # Rebuild filter choices when the active profile changes.
    original_reload_profile = App.reload_current_profile

    def reload_current_profile(self):
        for tab in getattr(self, "record_tabs", {}).values():
            if hasattr(tab, "_build_filters"):
                try:
                    tab._build_filters()
                except Exception:
                    pass
        original_reload_profile(self)
        for tab in getattr(self, "record_tabs", {}).values():
            if hasattr(tab, "_build_filters"):
                try:
                    tab._build_filters()
                    tab.refresh()
                except Exception:
                    pass

    App.reload_current_profile = reload_current_profile

    # Add Global Search button after normal startup.
    original_init = App.__init__

    def init(self, *args, **kwargs):
        original_init(self, *args, **kwargs)
        try:
            ttk.Button(
                self.profile_bar,
                text="Global Search",
                command=lambda: GlobalSearchDialog(self),
            ).pack(side="left", padx=3)
        except Exception:
            pass
        try:
            self.title(f"Forensic CV Manager Professional {PRO_VERSION}")
        except Exception:
            pass

    App.__init__ = init

    # Generate tab: keep Professional Tracking's granular export UI, then add presets.
    original_generate_tab = App._generate_tab

    def generate_tab(self):
        original_generate_tab(self)
        try:
            tabs = self.notebook.tabs()
            if tabs:
                frame = self.notebook.nametowidget(tabs[-1])
                add_preset_toolbar(self, frame)
        except Exception:
            pass

    App._generate_tab = generate_tab

    # Help menu additions while preserving the existing File/Tools/Help menus.
    original_menu = App._menu

    def menu(self):
        original_menu(self)
        try:
            menubar = self.nametowidget(self.cget("menu"))
            # Locate Help cascade by label.
            help_menu = None
            end = menubar.index("end")
            for i in range((end or -1) + 1):
                if menubar.entrycget(i, "label") == "Help":
                    menu_name = menubar.entrycget(i, "menu")
                    help_menu = self.nametowidget(menu_name)
                    break
            if help_menu is not None:
                help_menu.add_separator()
                help_menu.add_command(label="Professional Edition Changelog", command=self.show_changelog)
        except Exception:
            pass

    App._menu = menu
    App.show_about = professional_about
    App.show_changelog = show_changelog
    App.global_search = lambda self: GlobalSearchDialog(self)
    App.navigate_to_record = lambda self, table, row_id=None: navigate_to_record(self, table, row_id)
    App.generate_expert_report = lambda self: generate_expert_report_interactive(self)
    App.generate_annual_report = lambda self: generate_annual_report_interactive(self)

    _INSTALLED = True
