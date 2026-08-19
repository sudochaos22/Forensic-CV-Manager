from __future__ import annotations

"""
Professional Tracking extension for Forensic CV Manager.

Designed as a minimally invasive upgrade layer:
- Extends database tables/columns at runtime.
- Extends TABLE_CONFIG without replacing existing user customizations.
- Adds scrollable record dialogs for larger professional-history forms.
- Replaces the Generate CV tab with granular CV / metrics / appendix controls.
- Appends selected professional metrics and detailed appendices to Word/PDF.
- Adds selected-data ZIP/CSV export.
- Enhances the dashboard.

Sensitive case identifiers and notes are OFF by default in generated output.
"""

import csv
import html
import re
import tempfile
import zipfile
from collections import Counter
from datetime import date
from pathlib import Path
from typing import Any, Iterable

import tkinter as tk
from tkinter import filedialog, messagebox, ttk

import database as database_module
from date_utils import normalize_date


# ---------------------------------------------------------------------------
# Database definitions
# ---------------------------------------------------------------------------

CASEWORK_FIELDS = [
    "examination_date", "case_number", "requesting_agency", "case_type",
    "evidence_number", "device_type", "device_make_model", "device_size",
    "operating_system", "acquisition_type", "image_format",
    "data_acquired", "data_examined", "tools_used", "hours",
    "report_written", "testified", "status", "notes", "sort_order",
]

NEW_TABLE_FIELDS: dict[str, list[str]] = {
    "casework": CASEWORK_FIELDS,
    "forensic_reports": [
        "report_date", "case_number", "report_type", "role", "pages",
        "supplemental", "peer_reviewed", "subject_area", "notes", "sort_order",
    ],
    "peer_reviews": [
        "review_date", "case_number", "review_type", "role", "tools_methods",
        "hours", "changes_required", "outcome", "notes", "sort_order",
    ],
    "court_qualifications": [
        "qualification_date", "case_number", "court", "jurisdiction",
        "qualification_area", "ruling", "challenge_type", "notes", "sort_order",
    ],
    "tool_experience": [
        "tool_name", "category", "first_used", "last_used", "versions_used",
        "proficiency", "approx_examinations", "training_hours", "notes", "sort_order",
    ],
    "presentations": [
        "presentation_date", "title", "event", "organization", "location",
        "audience_size", "role", "hours", "notes", "sort_order",
    ],
    "publications": [
        "publication_date", "title", "publication", "publication_type",
        "authorship_role", "url", "citation", "notes", "sort_order",
    ],
    "validations": [
        "validation_date", "tool_name", "version", "validation_type", "dataset",
        "result", "role", "document_reference", "notes", "sort_order",
    ],
    "procedures": [
        "effective_date", "title", "organization", "document_type", "role",
        "status", "version", "notes", "sort_order",
    ],
    "mentoring": [
        "start_date", "end_date", "program", "mentee_group", "role",
        "hours", "outcome", "notes", "sort_order",
    ],
    "projects": [
        "start_date", "end_date", "title", "organization", "project_type",
        "role", "technologies", "impact", "notes", "sort_order",
    ],
    "professional_evidence": [
        "evidence_date", "evidence_type", "title", "related_area",
        "reference_location", "verification_url", "notes", "sort_order",
    ],
}

# Columns added to existing tables. SQLite is intentionally permissive here.
EXISTING_TABLE_COLUMNS: dict[str, dict[str, str]] = {
    "training": {
        "cpe_credits": "REAL",
        "related_certification": "TEXT",
        "delivery_method": "TEXT",
        "certificate_reference": "TEXT",
    },
    "certifications": {
        "renewal_date": "TEXT",
        "cpe_required": "REAL",
        "cpe_earned": "REAL",
        "verification_url": "TEXT",
    },
    "testimony": {
        "proceeding_type": "TEXT",
        "qualified_expert": "INTEGER DEFAULT 0",
        "qualification_area": "TEXT",
        "voir_dire": "INTEGER DEFAULT 0",
        "challenge_type": "TEXT",
        "testimony_hours": "REAL",
    },
    "teaching": {
        "students": "INTEGER",
        "curriculum_role": "TEXT",
        "evaluation_notes": "TEXT",
    },
    "skills": {
        "first_used": "TEXT",
        "last_used": "TEXT",
        "versions_used": "TEXT",
        "approx_examinations": "INTEGER",
    },
    "casework": {
        "image_format": "TEXT",
        "data_acquired": "TEXT",
        "data_examined": "TEXT",
    },
}

EXTENDED_DATE_FIELDS = {
    "examination_date", "report_date", "review_date", "qualification_date",
    "first_used", "last_used", "renewal_date", "presentation_date",
    "publication_date", "validation_date", "effective_date", "evidence_date",
}

ORDER_BY = {
    "casework": "examination_date DESC, sort_order, id DESC",
    "forensic_reports": "report_date DESC, sort_order, id DESC",
    "peer_reviews": "review_date DESC, sort_order, id DESC",
    "court_qualifications": "qualification_date DESC, sort_order, id DESC",
    "tool_experience": "last_used DESC, tool_name, sort_order, id DESC",
    "presentations": "presentation_date DESC, sort_order, id DESC",
    "publications": "publication_date DESC, sort_order, id DESC",
    "validations": "validation_date DESC, sort_order, id DESC",
    "procedures": "effective_date DESC, sort_order, id DESC",
    "mentoring": "start_date DESC, sort_order, id DESC",
    "projects": "start_date DESC, sort_order, id DESC",
    "professional_evidence": "evidence_date DESC, sort_order, id DESC",
}


def _sql_type(field: str) -> str:
    if field in {
        "hours", "training_hours", "cpe_credits", "cpe_required", "cpe_earned",
        "testimony_hours"
    }:
        return "REAL"
    if field in {
        "pages", "supplemental", "peer_reviewed", "changes_required",
        "audience_size", "approx_examinations", "students", "report_written",
        "testified", "qualified_expert", "voir_dire"
    }:
        return "INTEGER"
    if field == "sort_order":
        return "INTEGER DEFAULT 0"
    return "TEXT"


def _create_sql(table: str, fields: Iterable[str]) -> str:
    cols = [
        "id INTEGER PRIMARY KEY AUTOINCREMENT",
        "profile_id INTEGER NOT NULL",
    ]
    cols.extend(f"{field} {_sql_type(field)}" for field in fields)
    cols.append(
        "FOREIGN KEY(profile_id) REFERENCES profiles(id) ON DELETE CASCADE"
    )
    return f"CREATE TABLE IF NOT EXISTS {table} ({', '.join(cols)})"


def _append_unique(seq: list[str], values: Iterable[str]) -> None:
    for value in values:
        if value not in seq:
            seq.append(value)


_DB_INSTALLED = False


def install_database_extensions() -> None:
    """Extend database globals and wrap migration/list_rows once."""
    global _DB_INSTALLED
    if _DB_INSTALLED:
        return

    # Ensure Case Work exists even on an unmodified upstream checkout.
    for table, fields in NEW_TABLE_FIELDS.items():
        if table not in database_module.DATA_TABLES:
            database_module.DATA_TABLES.append(table)
        if table not in database_module.TABLE_FIELDS:
            database_module.TABLE_FIELDS[table] = list(fields)
        else:
            _append_unique(database_module.TABLE_FIELDS[table], fields)
        database_module.CREATE_DATA_TABLES.setdefault(
            table, _create_sql(table, database_module.TABLE_FIELDS[table])
        )

    # Extend existing table field allow-lists.
    for table, columns in EXISTING_TABLE_COLUMNS.items():
        if table not in database_module.TABLE_FIELDS:
            # Casework can be absent from original upstream.
            if table == "casework":
                database_module.TABLE_FIELDS[table] = list(CASEWORK_FIELDS)
            else:
                continue
        # Insert extra columns before sort_order where possible.
        fields = database_module.TABLE_FIELDS[table]
        for col in columns:
            if col in fields:
                continue
            if "sort_order" in fields:
                fields.insert(fields.index("sort_order"), col)
            else:
                fields.append(col)

    Database = database_module.Database
    original_migrate = Database._migrate_schema
    original_list_rows = Database.list_rows

    def extended_migrate(self):
        original_migrate(self)

        # Add any newly defined columns to existing tables without deleting data.
        all_columns: dict[str, dict[str, str]] = dict(EXISTING_TABLE_COLUMNS)
        # Also ensure every field in new/custom tables exists if an older custom
        # version of that table already exists.
        for table, fields in NEW_TABLE_FIELDS.items():
            all_columns.setdefault(table, {})
            for field in fields:
                if field == "sort_order":
                    all_columns[table].setdefault(field, "INTEGER DEFAULT 0")
                else:
                    all_columns[table].setdefault(field, _sql_type(field))

        for table, columns in all_columns.items():
            if not self._table_exists(table):
                continue
            existing = set(self._columns(table))
            for name, sql_type in columns.items():
                if name not in existing:
                    self.conn.execute(
                        f"ALTER TABLE {table} ADD COLUMN {name} {sql_type}"
                    )
        self.conn.commit()

    def extended_list_rows(self, table: str, order_by: str | None = None):
        if order_by is None and table in ORDER_BY:
            order_by = ORDER_BY[table]
        return original_list_rows(self, table, order_by)

    Database._migrate_schema = extended_migrate
    Database.list_rows = extended_list_rows
    _DB_INSTALLED = True


# ---------------------------------------------------------------------------
# UI table configuration
# ---------------------------------------------------------------------------

NEW_TABLE_CONFIG: dict[str, dict[str, Any]] = {
    "casework": {
        "label": "Case Work",
        "display": [
            "examination_date", "case_number", "case_type", "device_type",
            "device_size", "status",
        ],
        "fields": [
            ("examination_date", "Examination Date", "entry"),
            ("case_number", "Case Number", "entry"),
            ("requesting_agency", "Requesting Agency", "entry"),
            ("case_type", "Case Type", "combo", [
                "", "Cybercrime", "Child Exploitation / ICAC", "Homicide",
                "Fraud", "Narcotics", "Internal / Administrative", "Other",
            ]),
            ("evidence_number", "Evidence / Item Number", "entry"),
            ("device_type", "Device Type", "combo", [
                "", "Laptop", "Desktop", "Mobile Phone", "Tablet",
                "External HDD / SSD", "USB Flash Drive", "Memory Card",
                "Cloud Account", "Server", "Network Capture", "Memory / RAM",
                "Other",
            ]),
            ("device_make_model", "Device Make / Model", "entry"),
            ("device_size", "Device Capacity / Storage Size", "entry"),
            ("operating_system", "Operating System", "entry"),
            ("acquisition_type", "Acquisition Type", "combo", [
                "", "Physical", "Full File System", "File System",
                "Advanced Logical", "Logical", "Dead-box", "Live", "Triage",
                "Cloud / API", "Network Capture", "Memory Acquisition", "Other",
            ]),
            ("image_format", "Image / Acquisition Format", "entry"),
            ("data_acquired", "Data Acquired", "entry"),
            ("data_examined", "Data Examined", "entry"),
            ("tools_used", "Tools Used (semicolon separated)", "entry"),
            ("hours", "Examination Hours", "entry"),
            ("report_written", "Report Written", "check"),
            ("testified", "Testified", "check"),
            ("status", "Status", "combo", [
                "Complete", "Ongoing", "Pending", "Archived",
            ]),
            ("notes", "Notes / Work Performed", "text"),
        ],
    },
    "forensic_reports": {
        "label": "Forensic Reports",
        "display": ["report_date", "case_number", "report_type", "role", "pages"],
        "fields": [
            ("report_date", "Report Date", "entry"),
            ("case_number", "Case Number", "entry"),
            ("report_type", "Report Type", "combo", [
                "Forensic Examination", "Supplemental", "Technical",
                "Investigative", "Affidavit / Declaration", "Other",
            ]),
            ("role", "Role", "combo", [
                "Author", "Co-Author", "Technical Reviewer", "Approver", "Other",
            ]),
            ("pages", "Page Count", "entry"),
            ("supplemental", "Supplemental / Amended", "check"),
            ("peer_reviewed", "Peer Reviewed", "check"),
            ("subject_area", "Subject / Discipline", "entry"),
            ("notes", "Notes", "text"),
        ],
    },
    "peer_reviews": {
        "label": "Peer / Technical Reviews",
        "display": ["review_date", "case_number", "review_type", "role", "hours"],
        "fields": [
            ("review_date", "Review Date", "entry"),
            ("case_number", "Case Number / Reference", "entry"),
            ("review_type", "Review Type", "combo", [
                "Forensic Report", "Forensic Examination", "Acquisition",
                "Methodology", "Tool Output", "Validation", "Other",
            ]),
            ("role", "Role", "combo", [
                "Peer Reviewer", "Technical Reviewer", "Quality Reviewer",
                "Supervisor / Approver", "Other",
            ]),
            ("tools_methods", "Tools / Methods Reviewed", "entry"),
            ("hours", "Review Hours", "entry"),
            ("changes_required", "Changes Required", "check"),
            ("outcome", "Outcome", "entry"),
            ("notes", "Notes", "text"),
        ],
    },
    "court_qualifications": {
        "label": "Court Qualifications",
        "display": [
            "qualification_date", "court", "jurisdiction",
            "qualification_area", "ruling",
        ],
        "fields": [
            ("qualification_date", "Qualification Date", "entry"),
            ("case_number", "Case Number", "entry"),
            ("court", "Court", "entry"),
            ("jurisdiction", "Jurisdiction", "entry"),
            ("qualification_area", "Area of Expertise", "entry"),
            ("ruling", "Ruling / Result", "combo", [
                "Qualified", "Accepted Without Objection", "Limited",
                "Not Reached", "Other",
            ]),
            ("challenge_type", "Challenge / Voir Dire Notes", "entry"),
            ("notes", "Notes", "text"),
        ],
    },
    "tool_experience": {
        "label": "Tool Experience",
        "display": ["tool_name", "category", "first_used", "last_used", "proficiency"],
        "fields": [
            ("tool_name", "Tool / Platform", "entry"),
            ("category", "Category", "combo", [
                "Computer Forensics", "Mobile Forensics", "Cloud",
                "Memory Forensics", "Network Forensics", "Malware / RE",
                "eDiscovery", "Scripting / Automation", "Other",
            ]),
            ("first_used", "First Used", "entry"),
            ("last_used", "Last Used", "entry"),
            ("versions_used", "Versions Used", "entry"),
            ("proficiency", "Proficiency", "combo", [
                "", "Working Knowledge", "Proficient", "Advanced", "Expert / SME",
            ]),
            ("approx_examinations", "Approx. Examinations", "entry"),
            ("training_hours", "Related Training Hours", "entry"),
            ("notes", "Notes", "text"),
        ],
    },
    "presentations": {
        "label": "Presentations",
        "display": ["presentation_date", "title", "event", "organization", "role"],
        "fields": [
            ("presentation_date", "Presentation Date", "entry"),
            ("title", "Presentation Title", "entry"),
            ("event", "Conference / Event", "entry"),
            ("organization", "Organization", "entry"),
            ("location", "Location", "entry"),
            ("audience_size", "Approx. Audience Size", "entry"),
            ("role", "Role", "combo", [
                "Presenter", "Co-Presenter", "Panelist", "Moderator", "Briefing", "Other",
            ]),
            ("hours", "Presentation Hours", "entry"),
            ("notes", "Notes", "text"),
        ],
    },
    "publications": {
        "label": "Publications / Research",
        "display": ["publication_date", "title", "publication", "publication_type"],
        "fields": [
            ("publication_date", "Publication Date", "entry"),
            ("title", "Title", "entry"),
            ("publication", "Publisher / Venue", "entry"),
            ("publication_type", "Type", "combo", [
                "Article", "White Paper", "Research Paper", "Technical Note",
                "Book / Chapter", "Blog / Web", "Other",
            ]),
            ("authorship_role", "Authorship Role", "combo", [
                "Author", "Co-Author", "Contributor", "Reviewer", "Other",
            ]),
            ("url", "URL / DOI", "entry"),
            ("citation", "Citation", "entry"),
            ("notes", "Notes", "text"),
        ],
    },
    "validations": {
        "label": "Tool Validation / Testing",
        "display": ["validation_date", "tool_name", "version", "validation_type", "result"],
        "fields": [
            ("validation_date", "Validation Date", "entry"),
            ("tool_name", "Tool / Method", "entry"),
            ("version", "Version", "entry"),
            ("validation_type", "Validation Type", "combo", [
                "Initial Validation", "Version Verification", "Method Validation",
                "Regression / Re-Test", "Research Test", "Other",
            ]),
            ("dataset", "Test Dataset / Reference", "entry"),
            ("result", "Result", "combo", [
                "Passed", "Passed with Limitations", "Failed", "Research Only", "Other",
            ]),
            ("role", "Role", "entry"),
            ("document_reference", "Validation Document Reference", "entry"),
            ("notes", "Notes", "text"),
        ],
    },
    "procedures": {
        "label": "SOP / Policy Development",
        "display": ["effective_date", "title", "document_type", "role", "status"],
        "fields": [
            ("effective_date", "Effective / Revision Date", "entry"),
            ("title", "Document Title", "entry"),
            ("organization", "Organization", "entry"),
            ("document_type", "Document Type", "combo", [
                "SOP", "Policy", "Procedure", "Playbook", "Technical Guide",
                "Evidence Procedure", "Other",
            ]),
            ("role", "Role", "combo", [
                "Author", "Co-Author", "Lead Reviewer", "Contributor", "Approver", "Other",
            ]),
            ("status", "Status", "combo", [
                "Current", "Superseded", "Draft", "Retired", "Other",
            ]),
            ("version", "Version", "entry"),
            ("notes", "Notes", "text"),
        ],
    },
    "mentoring": {
        "label": "Mentoring / Supervision",
        "display": ["start_date", "end_date", "program", "role", "hours"],
        "fields": [
            ("start_date", "Start Date", "entry"),
            ("end_date", "End Date", "entry"),
            ("program", "Program / Activity", "entry"),
            ("mentee_group", "Mentee / Group Reference", "entry"),
            ("role", "Role", "combo", [
                "Mentor", "Coach", "Supervisor", "Technical Lead", "Peer Coach", "Other",
            ]),
            ("hours", "Hours", "entry"),
            ("outcome", "Outcome / Completion", "entry"),
            ("notes", "Notes", "text"),
        ],
    },
    "projects": {
        "label": "Major Projects",
        "display": ["start_date", "end_date", "title", "project_type", "role"],
        "fields": [
            ("start_date", "Start Date", "entry"),
            ("end_date", "End Date", "entry"),
            ("title", "Project / Accomplishment", "entry"),
            ("organization", "Organization", "entry"),
            ("project_type", "Project Type", "combo", [
                "Capability Development", "Lab / Infrastructure", "Automation",
                "Tool Deployment", "Research", "Major Investigation",
                "Program Development", "Other",
            ]),
            ("role", "Role", "entry"),
            ("technologies", "Technologies / Tools", "entry"),
            ("impact", "Impact / Result", "text"),
            ("notes", "Notes", "text"),
        ],
    },
    "professional_evidence": {
        "label": "Supporting Evidence",
        "display": ["evidence_date", "evidence_type", "title", "related_area"],
        "fields": [
            ("evidence_date", "Date", "entry"),
            ("evidence_type", "Evidence Type", "combo", [
                "Training Certificate", "Certification Certificate",
                "Transcript", "Court Record", "Award / Commendation",
                "Conference Agenda", "Publication", "Course Roster",
                "Validation Document", "Performance Record", "Other",
            ]),
            ("title", "Title / Description", "entry"),
            ("related_area", "Related CV Area", "entry"),
            ("reference_location", "File / Reference Location", "entry"),
            ("verification_url", "Verification URL", "entry"),
            ("notes", "Notes", "text"),
        ],
    },
}


EXISTING_CONFIG_ADDITIONS: dict[str, list[tuple]] = {
    "training": [
        ("cpe_credits", "CPE / CE Credits", "entry"),
        ("related_certification", "Related Certification", "entry"),
        ("delivery_method", "Delivery Method", "combo", [
            "", "Instructor-Led", "Self-Paced", "Virtual Live", "Conference", "Other",
        ]),
        ("certificate_reference", "Certificate / Evidence Reference", "entry"),
    ],
    "certifications": [
        ("renewal_date", "Renewal Date", "entry"),
        ("cpe_required", "CPE Required", "entry"),
        ("cpe_earned", "CPE Earned", "entry"),
        ("verification_url", "Verification URL", "entry"),
    ],
    "testimony": [
        ("proceeding_type", "Proceeding Type", "combo", [
            "", "Hearing", "Trial", "Deposition", "Grand Jury", "Motion Hearing", "Other",
        ]),
        ("qualified_expert", "Qualified / Accepted as Expert", "check"),
        ("qualification_area", "Expert Qualification Area", "entry"),
        ("voir_dire", "Voir Dire Conducted", "check"),
        ("challenge_type", "Daubert / Frye / Other Challenge", "entry"),
        ("testimony_hours", "Testimony / Proceeding Hours", "entry"),
    ],
    "teaching": [
        ("students", "Number of Students", "entry"),
        ("curriculum_role", "Curriculum Role", "combo", [
            "", "Developed", "Delivered", "Developed & Delivered", "Updated / Revised",
        ]),
        ("evaluation_notes", "Evaluation / Outcome Notes", "text"),
    ],
    "skills": [
        ("first_used", "First Used", "entry"),
        ("last_used", "Last Used", "entry"),
        ("versions_used", "Versions Used", "entry"),
        ("approx_examinations", "Approx. Examinations / Uses", "entry"),
    ],
    "casework": [
        ("image_format", "Image / Acquisition Format", "entry"),
        ("data_acquired", "Data Acquired", "entry"),
        ("data_examined", "Data Examined", "entry"),
    ],
}


def _field_names(config: dict[str, Any]) -> set[str]:
    return {field[0] for field in config.get("fields", [])}


def extend_table_config(
    table_config: dict[str, dict[str, Any]],
    date_fields: set[str],
    year_fields: set[str] | None = None,
) -> None:
    """Merge extension tabs/fields into the app's existing TABLE_CONFIG."""
    date_fields.update(EXTENDED_DATE_FIELDS)

    # Add missing tabs without replacing a user's customized Case Work config.
    for table, config in NEW_TABLE_CONFIG.items():
        table_config.setdefault(table, {
            "label": config["label"],
            "display": list(config["display"]),
            "fields": list(config["fields"]),
        })

    # Enrich existing forms.
    for table, additions in EXISTING_CONFIG_ADDITIONS.items():
        if table not in table_config:
            continue
        names = _field_names(table_config[table])
        for field in additions:
            if field[0] not in names:
                table_config[table]["fields"].append(field)
                names.add(field[0])


# ---------------------------------------------------------------------------
# Data / metrics helpers
# ---------------------------------------------------------------------------

def _rows(db, table: str) -> list[dict[str, Any]]:
    try:
        return db.list_rows(table)
    except Exception:
        return []


def _float(value: Any) -> float:
    try:
        return float(value or 0)
    except (TypeError, ValueError):
        return 0.0


def _int(value: Any) -> int:
    try:
        return int(float(value or 0))
    except (TypeError, ValueError):
        return 0


def _truthy(value: Any) -> bool:
    if isinstance(value, str):
        return value.strip().lower() in {"1", "yes", "true", "y", "checked"}
    return bool(value)


def _split_tools(value: Any) -> list[str]:
    text = str(value or "").strip()
    if not text:
        return []
    return [
        part.strip()
        for part in re.split(r"[;\n]+", text)
        if part.strip()
    ]


_SIZE_RE = re.compile(
    r"^\s*([\d,.]+)\s*(B|KB|MB|GB|TB|PB|KIB|MIB|GIB|TIB|PIB)?\s*$",
    re.I,
)


def parse_size_bytes(value: Any) -> float:
    text = str(value or "").strip().replace(",", "")
    if not text:
        return 0.0
    m = _SIZE_RE.match(text)
    if not m:
        return 0.0
    number = float(m.group(1))
    unit = (m.group(2) or "B").upper()
    powers = {
        "B": 0, "KB": 1, "KIB": 1, "MB": 2, "MIB": 2,
        "GB": 3, "GIB": 3, "TB": 4, "TIB": 4,
        "PB": 5, "PIB": 5,
    }
    return number * (1024 ** powers[unit])


def format_bytes(value: float) -> str:
    if value <= 0:
        return "0 B"
    units = ["B", "KB", "MB", "GB", "TB", "PB"]
    size = float(value)
    unit = units[0]
    for unit in units:
        if size < 1024 or unit == units[-1]:
            break
        size /= 1024
    return f"{size:,.2f} {unit}"


def professional_metrics(db) -> dict[str, Any]:
    casework = _rows(db, "casework")
    reports = _rows(db, "forensic_reports")
    reviews = _rows(db, "peer_reviews")
    testimony = _rows(db, "testimony")
    qualifications = _rows(db, "court_qualifications")
    training = _rows(db, "training")
    certifications = _rows(db, "certifications")
    teaching = _rows(db, "teaching")
    tools_history = _rows(db, "tool_experience")
    presentations = _rows(db, "presentations")
    publications = _rows(db, "publications")
    validations = _rows(db, "validations")
    procedures = _rows(db, "procedures")
    mentoring = _rows(db, "mentoring")
    projects = _rows(db, "projects")
    evidence = _rows(db, "professional_evidence")

    unique_cases = {
        str(r.get("case_number") or "").strip()
        for r in casework
        if str(r.get("case_number") or "").strip()
    }
    device_types = Counter(
        str(r.get("device_type") or "").strip()
        for r in casework if str(r.get("device_type") or "").strip()
    )
    case_types = Counter(
        str(r.get("case_type") or "").strip()
        for r in casework if str(r.get("case_type") or "").strip()
    )
    acquisition_types = Counter(
        str(r.get("acquisition_type") or "").strip()
        for r in casework if str(r.get("acquisition_type") or "").strip()
    )
    tools = Counter()
    for r in casework:
        tools.update(_split_tools(r.get("tools_used")))

    acquired_bytes = sum(parse_size_bytes(r.get("data_acquired")) for r in casework)
    examined_bytes = sum(parse_size_bytes(r.get("data_examined")) for r in casework)

    authored = sum(
        1 for r in reports
        if "author" in str(r.get("role") or "").lower()
        and "review" not in str(r.get("role") or "").lower()
    )
    expert = sum(1 for r in testimony if r.get("witness_type") == "Expert Witness")
    fact = sum(1 for r in testimony if r.get("witness_type") == "Fact Witness")
    qualified = sum(1 for r in testimony if _truthy(r.get("qualified_expert")))

    active_certs = sum(
        1 for r in certifications
        if str(r.get("status") or "").strip().lower() in {"active", "current", ""}
    )

    return {
        "casework": {
            "unique_cases": len(unique_cases),
            "examinations": len(casework),
            "hours": sum(_float(r.get("hours")) for r in casework),
            "reports_flagged": sum(1 for r in casework if _truthy(r.get("report_written"))),
            "testimony_flagged": sum(1 for r in casework if _truthy(r.get("testified"))),
            "device_types": device_types,
            "case_types": case_types,
            "acquisition_types": acquisition_types,
            "tools": tools,
            "data_acquired": acquired_bytes,
            "data_examined": examined_bytes,
        },
        "reports": {
            "total": len(reports),
            "authored": authored,
            "peer_reviewed": sum(1 for r in reports if _truthy(r.get("peer_reviewed"))),
            "supplemental": sum(1 for r in reports if _truthy(r.get("supplemental"))),
            "pages": sum(_int(r.get("pages")) for r in reports),
        },
        "reviews": {
            "total": len(reviews),
            "hours": sum(_float(r.get("hours")) for r in reviews),
            "changes_required": sum(
                1 for r in reviews if _truthy(r.get("changes_required"))
            ),
        },
        "court": {
            "testimony_total": len(testimony),
            "expert": expert,
            "fact": fact,
            "qualified_testimony": qualified,
            "qualification_records": len(qualifications),
            "hours": sum(_float(r.get("testimony_hours")) for r in testimony),
        },
        "training": {
            "records": len(training),
            "hours": sum(_float(r.get("hours")) for r in training),
            "cpe": sum(_float(r.get("cpe_credits")) for r in training),
            "certifications": len(certifications),
            "active_certifications": active_certs,
        },
        "teaching": {
            "records": len(teaching),
            "hours": sum(_float(r.get("hours")) for r in teaching),
            "students": sum(_int(r.get("students")) for r in teaching),
        },
        "tools": {
            "tracked": len(tools_history),
            "casework_tools": tools,
        },
        "professional_output": {
            "presentations": len(presentations),
            "publications": len(publications),
        },
        "quality": {
            "validations": len(validations),
            "procedures": len(procedures),
            "peer_reviews": len(reviews),
        },
        "leadership": {
            "mentoring": len(mentoring),
            "mentoring_hours": sum(_float(r.get("hours")) for r in mentoring),
            "projects": len(projects),
            "evidence_records": len(evidence),
        },
    }


# ---------------------------------------------------------------------------
# Export options
# ---------------------------------------------------------------------------

CV_OPTIONS = [
    ("summary", "Professional Summary", 1),
    ("core_training", "Core Training", 1),
    ("employment", "Work Experience", 1),
    ("teaching", "Teaching Experience", 1),
    ("organizations", "Professional Organizations", 1),
    ("certifications", "Certifications", 1),
    ("skills", "Skills and Tools", 1),
    ("education", "Education", 1),
    ("testimony", "Courtroom Testimony", 1),
    ("achievements", "Professional Achievements", 1),
    ("full_training", "Detailed Training Appendix (legacy)", 0),
]

SUMMARY_OPTIONS = [
    ("forensic_case_summary", "Forensic Case Work Summary", 1),
    ("forensic_report_summary", "Forensic Report Summary", 1),
    ("court_summary", "Court / Expert Qualification Summary", 1),
    ("training_summary", "Training / CPE Summary", 1),
    ("tool_summary", "Forensic Tool Usage Summary", 1),
    ("teaching_summary", "Teaching / Instruction Summary", 0),
    ("quality_summary", "Peer Review / Validation / SOP Summary", 0),
    ("publication_summary", "Presentations / Publications Summary", 0),
    ("leadership_summary", "Mentoring / Major Projects Summary", 0),
]

DETAIL_OPTIONS = [
    ("detail_casework", "Case Work", 0),
    ("detail_reports", "Forensic Reports", 0),
    ("detail_peer_reviews", "Peer / Technical Reviews", 0),
    ("detail_court_qualifications", "Court Qualifications", 0),
    ("detail_testimony", "Courtroom Testimony", 0),
    ("detail_tool_experience", "Tool Experience", 0),
    ("detail_training", "Training / CPE", 0),
    ("detail_certifications", "Certifications / Maintenance", 0),
    ("detail_teaching", "Teaching", 0),
    ("detail_presentations", "Presentations", 0),
    ("detail_publications", "Publications / Research", 0),
    ("detail_validations", "Tool Validation / Testing", 0),
    ("detail_procedures", "SOP / Policy Development", 0),
    ("detail_mentoring", "Mentoring / Supervision", 0),
    ("detail_projects", "Major Projects", 0),
    ("detail_evidence", "Supporting Evidence Index", 0),
]

PRIVACY_OPTIONS = [
    ("include_case_numbers", "Include case numbers", 0),
    ("include_evidence_numbers", "Include evidence / item numbers", 0),
    ("include_notes", "Include free-text notes", 0),
    ("include_verification_links", "Include verification URLs / references", 0),
]


DETAIL_TABLES = {
    "detail_casework": (
        "casework", "Detailed Case Work",
        [
            ("examination_date", "Date"),
            ("case_number", "Case #"),
            ("evidence_number", "Item #"),
            ("case_type", "Case Type"),
            ("device_type", "Device"),
            ("device_make_model", "Make / Model"),
            ("device_size", "Capacity"),
            ("acquisition_type", "Acquisition"),
            ("image_format", "Format"),
            ("data_acquired", "Acquired"),
            ("data_examined", "Examined"),
            ("tools_used", "Tools"),
            ("hours", "Hours"),
            ("status", "Status"),
            ("notes", "Notes"),
        ],
    ),
    "detail_reports": (
        "forensic_reports", "Detailed Forensic Reports",
        [
            ("report_date", "Date"), ("case_number", "Case #"),
            ("report_type", "Type"), ("role", "Role"), ("pages", "Pages"),
            ("peer_reviewed", "Peer Reviewed"), ("subject_area", "Subject"),
            ("notes", "Notes"),
        ],
    ),
    "detail_peer_reviews": (
        "peer_reviews", "Detailed Peer / Technical Reviews",
        [
            ("review_date", "Date"), ("case_number", "Case / Ref"),
            ("review_type", "Review"), ("role", "Role"),
            ("tools_methods", "Tools / Methods"), ("hours", "Hours"),
            ("changes_required", "Changes"), ("outcome", "Outcome"),
            ("notes", "Notes"),
        ],
    ),
    "detail_court_qualifications": (
        "court_qualifications", "Detailed Court Qualifications",
        [
            ("qualification_date", "Date"), ("case_number", "Case #"),
            ("court", "Court"), ("jurisdiction", "Jurisdiction"),
            ("qualification_area", "Qualification Area"), ("ruling", "Ruling"),
            ("challenge_type", "Challenge"), ("notes", "Notes"),
        ],
    ),
    "detail_testimony": (
        "testimony", "Detailed Courtroom Testimony",
        [
            ("testimony_date", "Date"), ("case_number", "Case #"),
            ("court", "Court"), ("jurisdiction", "Jurisdiction"),
            ("witness_type", "Witness"), ("proceeding_type", "Proceeding"),
            ("qualification_area", "Qualification"), ("testimony_hours", "Hours"),
            ("notes", "Notes"),
        ],
    ),
    "detail_tool_experience": (
        "tool_experience", "Detailed Tool Experience",
        [
            ("tool_name", "Tool"), ("category", "Category"),
            ("first_used", "First Used"), ("last_used", "Last Used"),
            ("versions_used", "Versions"), ("proficiency", "Proficiency"),
            ("approx_examinations", "Approx. Exams"),
            ("training_hours", "Training Hrs"), ("notes", "Notes"),
        ],
    ),
    "detail_training": (
        "training", "Detailed Training / CPE",
        [
            ("attended_date", "Date"), ("course_name", "Course"),
            ("provider", "Provider"), ("hours", "Hours"),
            ("cpe_credits", "CPE"), ("related_certification", "Related Cert"),
            ("delivery_method", "Delivery"), ("certificate_reference", "Evidence"),
            ("notes", "Notes"),
        ],
    ),
    "detail_certifications": (
        "certifications", "Detailed Certifications / Maintenance",
        [
            ("certification", "Certification"),
            ("issuing_organization", "Issuer"), ("earned_date", "Earned"),
            ("expiration_date", "Expires"), ("renewal_date", "Renewed"),
            ("status", "Status"), ("cpe_required", "CPE Req."),
            ("cpe_earned", "CPE Earned"), ("credential_number", "Credential"),
            ("verification_url", "Verification"), ("notes", "Notes"),
        ],
    ),
    "detail_teaching": (
        "teaching", "Detailed Teaching / Instruction",
        [
            ("start_date", "Start"), ("end_date", "End"),
            ("organization", "Organization"), ("role", "Role"),
            ("course_name", "Course"), ("hours", "Hours"),
            ("students", "Students"), ("curriculum_role", "Curriculum"),
            ("description", "Description"), ("evaluation_notes", "Evaluation"),
        ],
    ),
    "detail_presentations": (
        "presentations", "Detailed Presentations",
        [
            ("presentation_date", "Date"), ("title", "Title"),
            ("event", "Event"), ("organization", "Organization"),
            ("role", "Role"), ("audience_size", "Audience"),
            ("hours", "Hours"), ("notes", "Notes"),
        ],
    ),
    "detail_publications": (
        "publications", "Detailed Publications / Research",
        [
            ("publication_date", "Date"), ("title", "Title"),
            ("publication", "Publisher / Venue"), ("publication_type", "Type"),
            ("authorship_role", "Role"), ("citation", "Citation"),
            ("url", "URL"), ("notes", "Notes"),
        ],
    ),
    "detail_validations": (
        "validations", "Detailed Tool Validation / Testing",
        [
            ("validation_date", "Date"), ("tool_name", "Tool"),
            ("version", "Version"), ("validation_type", "Type"),
            ("dataset", "Dataset"), ("result", "Result"), ("role", "Role"),
            ("document_reference", "Reference"), ("notes", "Notes"),
        ],
    ),
    "detail_procedures": (
        "procedures", "Detailed SOP / Policy Development",
        [
            ("effective_date", "Date"), ("title", "Title"),
            ("organization", "Organization"), ("document_type", "Type"),
            ("role", "Role"), ("status", "Status"), ("version", "Version"),
            ("notes", "Notes"),
        ],
    ),
    "detail_mentoring": (
        "mentoring", "Detailed Mentoring / Supervision",
        [
            ("start_date", "Start"), ("end_date", "End"),
            ("program", "Program"), ("mentee_group", "Mentee / Group"),
            ("role", "Role"), ("hours", "Hours"), ("outcome", "Outcome"),
            ("notes", "Notes"),
        ],
    ),
    "detail_projects": (
        "projects", "Detailed Major Projects",
        [
            ("start_date", "Start"), ("end_date", "End"), ("title", "Project"),
            ("organization", "Organization"), ("project_type", "Type"),
            ("role", "Role"), ("technologies", "Technologies"),
            ("impact", "Impact"), ("notes", "Notes"),
        ],
    ),
    "detail_evidence": (
        "professional_evidence", "Supporting Evidence Index",
        [
            ("evidence_date", "Date"), ("evidence_type", "Type"),
            ("title", "Title"), ("related_area", "Related Area"),
            ("reference_location", "Reference Location"),
            ("verification_url", "Verification"), ("notes", "Notes"),
        ],
    ),
}


def _filtered_fields(option_key: str, fields: list[tuple[str, str]], options: dict[str, Any]):
    result = []
    for field, label in fields:
        if field == "case_number" and not options.get("include_case_numbers", False):
            continue
        if field == "evidence_number" and not options.get("include_evidence_numbers", False):
            continue
        if field == "notes" and not options.get("include_notes", False):
            continue
        if field in {
            "verification_url", "reference_location", "document_reference",
            "certificate_reference", "credential_number", "url"
        } and not options.get("include_verification_links", False):
            continue
        result.append((field, label))
    return result


# ---------------------------------------------------------------------------
# Word appendices
# ---------------------------------------------------------------------------

def _word_heading(doc, text: str, level: int = 1):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper() if level == 1 else text)
    r.bold = True
    r.font.size = Pt(11 if level == 1 else 9.5)
    if level == 1:
        r.font.color.rgb = RGBColor(31, 78, 121)
    return p


def _word_bullets(doc, lines: Iterable[str]):
    from docx.shared import Pt
    for line in lines:
        if not str(line).strip():
            continue
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.add_run(str(line))


def _summary_sections(db, options: dict[str, Any]) -> list[tuple[str, list[str]]]:
    m = professional_metrics(db)
    sections: list[tuple[str, list[str]]] = []

    if options.get("forensic_case_summary"):
        c = m["casework"]
        lines = [
            f"Unique cases documented: {c['unique_cases']:,}",
            f"Forensic examinations / devices documented: {c['examinations']:,}",
            f"Documented examination hours: {c['hours']:,.1f}",
            f"Examinations with reports: {c['reports_flagged']:,}",
            f"Examinations associated with testimony: {c['testimony_flagged']:,}",
        ]
        if c["data_acquired"]:
            lines.append(f"Documented data acquired: {format_bytes(c['data_acquired'])}")
        if c["data_examined"]:
            lines.append(f"Documented data examined: {format_bytes(c['data_examined'])}")
        if c["device_types"]:
            lines.append(
                "Device mix: " + ", ".join(
                    f"{k} ({v})" for k, v in c["device_types"].most_common()
                )
            )
        if c["acquisition_types"]:
            lines.append(
                "Acquisition methods: " + ", ".join(
                    f"{k} ({v})" for k, v in c["acquisition_types"].most_common()
                )
            )
        sections.append(("Digital Forensic Case Experience", lines))

    if options.get("forensic_report_summary"):
        r = m["reports"]
        sections.append(("Forensic Reports", [
            f"Reports tracked: {r['total']:,}",
            f"Reports authored / co-authored: {r['authored']:,}",
            f"Peer-reviewed reports: {r['peer_reviewed']:,}",
            f"Supplemental / amended reports: {r['supplemental']:,}",
            f"Documented report pages: {r['pages']:,}",
        ]))

    if options.get("court_summary"):
        c = m["court"]
        sections.append(("Court and Expert Qualification Experience", [
            f"Testimony appearances tracked: {c['testimony_total']:,}",
            f"Expert-witness appearances: {c['expert']:,}",
            f"Fact-witness appearances: {c['fact']:,}",
            f"Testimony records marked qualified as expert: {c['qualified_testimony']:,}",
            f"Separate court qualification records: {c['qualification_records']:,}",
            f"Documented testimony / proceeding hours: {c['hours']:,.1f}",
        ]))

    if options.get("training_summary"):
        t = m["training"]
        sections.append(("Professional Development", [
            f"Training records: {t['records']:,}",
            f"Documented training hours: {t['hours']:,.1f}",
            f"Documented CPE / CE credits: {t['cpe']:,.1f}",
            f"Certifications tracked: {t['certifications']:,}",
            f"Active/current certifications: {t['active_certifications']:,}",
        ]))

    if options.get("tool_summary"):
        t = m["tools"]
        lines = [f"Tools/platform histories tracked: {t['tracked']:,}"]
        if t["casework_tools"]:
            lines.append(
                "Most-used tools in documented casework: " + ", ".join(
                    f"{k} ({v})" for k, v in t["casework_tools"].most_common(12)
                )
            )
        sections.append(("Forensic Tools and Platforms", lines))

    if options.get("teaching_summary"):
        t = m["teaching"]
        sections.append(("Teaching and Instruction", [
            f"Teaching records: {t['records']:,}",
            f"Documented instructional hours: {t['hours']:,.1f}",
            f"Documented students / attendees: {t['students']:,}",
        ]))

    if options.get("quality_summary"):
        q = m["quality"]
        sections.append(("Quality Assurance and Methodology", [
            f"Peer / technical reviews: {q['peer_reviews']:,}",
            f"Tool validation / testing records: {q['validations']:,}",
            f"SOP / policy / procedure records: {q['procedures']:,}",
        ]))

    if options.get("publication_summary"):
        p = m["professional_output"]
        sections.append(("Presentations and Publications", [
            f"Presentations / briefings tracked: {p['presentations']:,}",
            f"Publications / research records: {p['publications']:,}",
        ]))

    if options.get("leadership_summary"):
        l = m["leadership"]
        sections.append(("Leadership, Mentoring and Major Projects", [
            f"Mentoring / supervision records: {l['mentoring']:,}",
            f"Documented mentoring / supervision hours: {l['mentoring_hours']:,.1f}",
            f"Major projects / accomplishments tracked: {l['projects']:,}",
            f"Supporting-evidence records indexed: {l['evidence_records']:,}",
        ]))

    return sections


def _has_detail(options: dict[str, Any]) -> bool:
    return any(options.get(key) for key, _, _ in DETAIL_OPTIONS)


def append_professional_word(db, path: str | Path, options: dict[str, Any]) -> Path:
    from docx import Document
    from docx.enum.section import WD_ORIENT, WD_SECTION
    from docx.shared import Inches, Pt

    path = Path(path)
    doc = Document(path)

    for heading, lines in _summary_sections(db, options):
        _word_heading(doc, heading)
        _word_bullets(doc, lines)

    if _has_detail(options):
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = Inches(0.35)
        section.right_margin = Inches(0.35)
        _word_heading(doc, "Detailed Professional Record Appendices")

        for key, _, _ in DETAIL_OPTIONS:
            if not options.get(key):
                continue
            table_name, title, raw_fields = DETAIL_TABLES[key]
            rows = _rows(db, table_name)
            fields = _filtered_fields(key, raw_fields, options)
            if not rows or not fields:
                continue

            _word_heading(doc, title, level=2)
            table = doc.add_table(rows=1, cols=len(fields))
            table.style = "Table Grid"
            for i, (_, label) in enumerate(fields):
                cell = table.rows[0].cells[i]
                cell.text = label
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(7)

            for row in rows:
                cells = table.add_row().cells
                for i, (field, _) in enumerate(fields):
                    value = row.get(field, "")
                    if field in {
                        "supplemental", "peer_reviewed", "changes_required",
                        "report_written", "testified", "qualified_expert", "voir_dire"
                    }:
                        value = "Yes" if _truthy(value) else "No"
                    cells[i].text = "" if value is None else str(value)
                    for p in cells[i].paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(6.5)
            doc.add_paragraph()

    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# PDF appendices
# ---------------------------------------------------------------------------

def _generate_appendix_pdf(db, path: Path, options: dict[str, Any]) -> bool:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import LETTER, landscape
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        LongTable, Paragraph, SimpleDocTemplate, Spacer, TableStyle,
    )

    summaries = _summary_sections(db, options)
    has_detail = _has_detail(options)
    if not summaries and not has_detail:
        return False

    pagesize = landscape(LETTER) if has_detail else LETTER
    doc = SimpleDocTemplate(
        str(path), pagesize=pagesize,
        leftMargin=.35 * inch, rightMargin=.35 * inch,
        topMargin=.4 * inch, bottomMargin=.4 * inch,
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle(
        "ExtH1", parent=styles["Heading1"], fontSize=12, leading=14,
        textColor=colors.HexColor("#1F4E79"), spaceBefore=8, spaceAfter=5,
    )
    h2 = ParagraphStyle(
        "ExtH2", parent=styles["Heading2"], fontSize=10, leading=12,
        spaceBefore=7, spaceAfter=4,
    )
    body = ParagraphStyle(
        "ExtBody", parent=styles["BodyText"], fontSize=8.2, leading=10,
        alignment=TA_LEFT, spaceAfter=2,
    )
    cell_style = ParagraphStyle(
        "ExtCell", parent=styles["BodyText"], fontSize=5.8, leading=7,
    )
    head_style = ParagraphStyle(
        "ExtHead", parent=cell_style, fontName="Helvetica-Bold",
    )

    story = []
    for heading, lines in summaries:
        story.append(Paragraph(html.escape(heading), h1))
        for line in lines:
            story.append(Paragraph("• " + html.escape(str(line)), body))
        story.append(Spacer(1, 4))

    if has_detail:
        story.append(Paragraph("Detailed Professional Record Appendices", h1))
        for key, _, _ in DETAIL_OPTIONS:
            if not options.get(key):
                continue
            table_name, title, raw_fields = DETAIL_TABLES[key]
            rows = _rows(db, table_name)
            fields = _filtered_fields(key, raw_fields, options)
            if not rows or not fields:
                continue
            story.append(Paragraph(html.escape(title), h2))
            data = [[Paragraph(html.escape(label), head_style) for _, label in fields]]
            for row in rows:
                out_row = []
                for field, _ in fields:
                    value = row.get(field, "")
                    if field in {
                        "supplemental", "peer_reviewed", "changes_required",
                        "report_written", "testified", "qualified_expert", "voir_dire"
                    }:
                        value = "Yes" if _truthy(value) else "No"
                    out_row.append(
                        Paragraph(html.escape("" if value is None else str(value)), cell_style)
                    )
                data.append(out_row)

            table = LongTable(data, repeatRows=1, hAlign="LEFT")
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9EAF7")),
                ("GRID", (0, 0), (-1, -1), .25, colors.HexColor("#A0A0A0")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 2),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]))
            story.append(table)
            story.append(Spacer(1, 7))

    doc.build(story)
    return True


def _merge_pdfs(base: Path, appendix: Path, output: Path) -> None:
    try:
        import pymupdf
    except ImportError:
        import fitz as pymupdf  # compatibility fallback

    result = pymupdf.open()
    base_doc = pymupdf.open(str(base))
    result.insert_pdf(base_doc)
    base_doc.close()
    appendix_doc = pymupdf.open(str(appendix))
    result.insert_pdf(appendix_doc)
    appendix_doc.close()
    result.save(str(output))
    result.close()


def generate_extended_docx(db, output_path: str | Path, options: dict[str, Any]) -> Path:
    from cv_generator import generate_cv
    output = Path(output_path)
    base_options = dict(options)
    # Suppress the older custom Case Work block if a prior local build added it.
    base_options["casework_summary"] = False
    generate_cv(db, output, base_options)
    append_professional_word(db, output, options)
    return output


def generate_extended_pdf(
    db, output_path: str | Path, options: dict[str, Any],
    *, theme_name: str = "Professional"
) -> Path:
    from pdf_export import generate_pdf
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    base_options = dict(options)
    base_options["casework_summary"] = False

    with tempfile.TemporaryDirectory(prefix="fcv_professional_") as td:
        td = Path(td)
        base_pdf = td / "base.pdf"
        appendix_pdf = td / "appendix.pdf"
        generate_pdf(db, base_pdf, base_options, theme_name=theme_name)
        if _generate_appendix_pdf(db, appendix_pdf, options):
            _merge_pdfs(base_pdf, appendix_pdf, output)
        else:
            output.write_bytes(base_pdf.read_bytes())
    return output


# ---------------------------------------------------------------------------
# CSV / ZIP selective export
# ---------------------------------------------------------------------------

def export_selected_data_zip(db, output_path: str | Path, options: dict[str, Any]) -> Path:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    selections = [
        (key, DETAIL_TABLES[key])
        for key, _, _ in DETAIL_OPTIONS
        if options.get(key)
    ]
    if not selections:
        raise ValueError(
            "Select at least one item under Detailed Appendices before exporting data."
        )

    with tempfile.TemporaryDirectory(prefix="fcv_csv_") as td:
        td = Path(td)
        manifest = {
            "format": "Forensic CV Manager Selected Professional Data",
            "generated": date.today().isoformat(),
            "privacy": {
                k: bool(options.get(k))
                for k, _, _ in PRIVACY_OPTIONS
            },
            "files": [],
        }

        for key, (table_name, title, raw_fields) in selections:
            rows = _rows(db, table_name)
            fields = _filtered_fields(key, raw_fields, options)
            if not fields:
                continue

            file_name = re.sub(r"[^A-Za-z0-9_-]+", "_", table_name) + ".csv"
            csv_path = td / file_name
            with csv_path.open("w", newline="", encoding="utf-8-sig") as f:
                writer = csv.writer(f)
                writer.writerow([label for _, label in fields])
                for row in rows:
                    values = []
                    for field, _ in fields:
                        value = row.get(field, "")
                        if field in {
                            "supplemental", "peer_reviewed", "changes_required",
                            "report_written", "testified", "qualified_expert", "voir_dire"
                        }:
                            value = "Yes" if _truthy(value) else "No"
                        values.append("" if value is None else value)
                    writer.writerow(values)
            manifest["files"].append({"file": file_name, "section": title, "rows": len(rows)})

        manifest_path = td / "manifest.json"
        manifest_path.write_text(json_dumps(manifest), encoding="utf-8")

        with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as z:
            for file in td.iterdir():
                z.write(file, file.name)
    return output


def json_dumps(value: Any) -> str:
    import json
    return json.dumps(value, indent=2, ensure_ascii=False)


# ---------------------------------------------------------------------------
# Scrollable record dialog
# ---------------------------------------------------------------------------

NUMERIC_FLOAT_FIELDS = {
    "hours", "training_hours", "cpe_credits", "cpe_required", "cpe_earned",
    "testimony_hours",
}
NUMERIC_INT_FIELDS = {
    "pages", "audience_size", "approx_examinations", "students",
}


class ScrollableRecordDialog(tk.Toplevel):
    def __init__(self, parent, table: str, config: dict[str, Any], initial=None):
        super().__init__(parent)
        self.title(("Edit " if initial else "Add ") + config["label"])
        self.transient(parent)
        self.grab_set()
        self.resizable(True, True)
        self.result = None
        self.vars: dict[str, Any] = {}
        self.widgets: dict[str, Any] = {}
        self._date_fields = set(EXTENDED_DATE_FIELDS) | {
            "start_date", "end_date", "graduation_date", "attended_date",
            "expiration_date", "earned_date", "testimony_date", "achievement_date",
        }
        self._year_fields = {"start_year", "end_year"}

        outer = ttk.Frame(self)
        outer.pack(fill="both", expand=True)
        canvas = tk.Canvas(outer, highlightthickness=0)
        scroll = ttk.Scrollbar(outer, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=scroll.set)
        scroll.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)

        form = ttk.Frame(canvas, padding=10)
        window = canvas.create_window((0, 0), window=form, anchor="nw")
        form.columnconfigure(1, weight=1)

        def sync_region(_event=None):
            canvas.configure(scrollregion=canvas.bbox("all"))
        def sync_width(event):
            canvas.itemconfigure(window, width=event.width)
        form.bind("<Configure>", sync_region)
        canvas.bind("<Configure>", sync_width)

        row = 0
        for field in config["fields"]:
            name, label, kind, *extra = field
            ttk.Label(form, text=label + ":").grid(
                row=row, column=0, sticky="nw", padx=8, pady=5
            )
            value = "" if initial is None or initial.get(name) is None else initial.get(name)

            if kind == "text":
                widget = tk.Text(form, width=65, height=4, wrap="word")
                widget.insert("1.0", str(value))
                widget.grid(row=row, column=1, sticky="nsew", padx=8, pady=5)
            elif kind == "check":
                var = tk.IntVar(value=int(value or 0))
                widget = ttk.Checkbutton(form, variable=var)
                widget.grid(row=row, column=1, sticky="w", padx=8, pady=5)
                self.vars[name] = var
            elif kind == "combo":
                var = tk.StringVar(value=str(value))
                vals = extra[0] if extra else []
                widget = ttk.Combobox(
                    form, textvariable=var, values=vals, state="readonly"
                )
                if not value and vals:
                    var.set(vals[0])
                widget.grid(row=row, column=1, sticky="ew", padx=8, pady=5)
                self.vars[name] = var
            else:
                var = tk.StringVar(value=str(value))
                widget = ttk.Entry(form, textvariable=var, width=65)
                widget.grid(row=row, column=1, sticky="ew", padx=8, pady=5)
                self.vars[name] = var
            self.widgets[name] = widget
            row += 1

        buttons = ttk.Frame(form)
        buttons.grid(row=row, column=0, columnspan=2, sticky="e", padx=8, pady=10)
        ttk.Button(buttons, text="Cancel", command=self.destroy).pack(
            side="right", padx=4
        )
        ttk.Button(buttons, text="Save", command=self.save).pack(
            side="right", padx=4
        )

        screen_h = max(self.winfo_screenheight(), 700)
        self.geometry(f"860x{min(780, screen_h - 100)}")
        self.bind("<Escape>", lambda e: self.destroy())

        def wheel(event):
            delta = -1 if event.delta > 0 else 1
            canvas.yview_scroll(delta * 3, "units")
        canvas.bind_all("<MouseWheel>", wheel)
        self.protocol("WM_DELETE_WINDOW", self.destroy)

    def destroy(self):
        try:
            self.unbind_all("<MouseWheel>")
        except Exception:
            pass
        super().destroy()

    def save(self):
        result = {}
        for field, widget in self.widgets.items():
            if isinstance(widget, tk.Text):
                value = widget.get("1.0", "end").strip()
            else:
                value = self.vars[field].get()

            if field in NUMERIC_FLOAT_FIELDS:
                try:
                    value = float(value) if str(value).strip() else None
                except ValueError:
                    messagebox.showerror(
                        "Invalid Number", f"{field.replace('_', ' ').title()} must be a number.",
                        parent=self,
                    )
                    return
            elif field in NUMERIC_INT_FIELDS:
                try:
                    value = int(float(value)) if str(value).strip() else None
                except ValueError:
                    messagebox.showerror(
                        "Invalid Number", f"{field.replace('_', ' ').title()} must be a whole number.",
                        parent=self,
                    )
                    return
            elif field in self._date_fields:
                try:
                    value = normalize_date(value, allow_present=field == "end_date")
                except ValueError as exc:
                    messagebox.showerror("Invalid Date", str(exc), parent=self)
                    widget.focus_set()
                    return
            elif field in self._year_fields:
                try:
                    value = normalize_date(
                        value, allow_present=field == "end_year", allow_year_only=True
                    )
                except ValueError as exc:
                    messagebox.showerror("Invalid Year", str(exc), parent=self)
                    widget.focus_set()
                    return
            result[field] = value

        self.result = result
        self.destroy()


# ---------------------------------------------------------------------------
# Enhanced App methods
# ---------------------------------------------------------------------------

def _add_check_grid(parent, app, items, columns=2):
    for i, (key, label, default) in enumerate(items):
        var = tk.IntVar(value=default)
        ttk.Checkbutton(parent, text=label, variable=var).grid(
            row=i // columns, column=i % columns,
            sticky="w", padx=10, pady=4,
        )
        app.cv_options[key] = var


def enhanced_generate_tab(self):
    tab = ttk.Frame(self.notebook)
    self.notebook.add(tab, text="Generate CV")

    header = ttk.Frame(tab, padding=(16, 14, 16, 4))
    header.pack(fill="x")
    ttk.Label(
        header, text="Create CV / Professional Record Output",
        style="Title.TLabel"
    ).pack(anchor="w")
    ttk.Label(
        header,
        text=(
            "Choose exactly what is exported. Aggregate metrics are suitable for "
            "CV/voir-dire use; detailed appendices and identifiers are opt-in."
        ),
    ).pack(anchor="w", pady=(4, 0))

    book = ttk.Notebook(tab)
    book.pack(fill="both", expand=True, padx=16, pady=8)
    self.cv_options = {}

    cv_tab = ttk.Frame(book, padding=12)
    metrics_tab = ttk.Frame(book, padding=12)
    details_tab = ttk.Frame(book, padding=12)
    book.add(cv_tab, text="CV Sections")
    book.add(metrics_tab, text="Professional Metrics")
    book.add(details_tab, text="Detailed Appendices / Privacy")

    ttk.Label(
        cv_tab,
        text="Traditional CV sections",
        style="Title.TLabel"
    ).grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 8))
    cv_checks = ttk.Frame(cv_tab)
    cv_checks.grid(row=1, column=0, columnspan=2, sticky="nw")
    _add_check_grid(cv_checks, self, CV_OPTIONS, columns=2)

    ttk.Label(
        metrics_tab,
        text="Aggregate professional experience metrics (no case identifiers)",
        style="Title.TLabel"
    ).grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 8))
    metrics_checks = ttk.Frame(metrics_tab)
    metrics_checks.grid(row=1, column=0, columnspan=2, sticky="nw")
    _add_check_grid(metrics_checks, self, SUMMARY_OPTIONS, columns=2)

    detail_box = ttk.LabelFrame(details_tab, text="Detailed Appendices", padding=8)
    detail_box.pack(fill="x", pady=(0, 10))
    _add_check_grid(detail_box, self, DETAIL_OPTIONS, columns=2)

    privacy_box = ttk.LabelFrame(
        details_tab,
        text="Sensitive / Identifying Fields — OFF by default",
        padding=8
    )
    privacy_box.pack(fill="x", pady=(0, 10))
    _add_check_grid(privacy_box, self, PRIVACY_OPTIONS, columns=2)

    control = ttk.Frame(details_tab)
    control.pack(fill="x")
    detail_keys = [key for key, _, _ in DETAIL_OPTIONS]

    def set_details(value: int):
        for key in detail_keys:
            self.cv_options[key].set(value)

    ttk.Button(
        control, text="Select All Detailed Appendices",
        command=lambda: set_details(1)
    ).pack(side="left", padx=(0, 6))
    ttk.Button(
        control, text="Clear Detailed Appendices",
        command=lambda: set_details(0)
    ).pack(side="left")

    ttk.Label(
        details_tab,
        text=(
            "Tip: You can export detailed records without case numbers or notes. "
            "Enable identifiers only for an authorized need-to-know output."
        ),
        wraplength=900,
    ).pack(anchor="w", pady=(12, 0))

    actions = ttk.Frame(tab, padding=(16, 4, 16, 14))
    actions.pack(fill="x")
    ttk.Button(
        actions, text="Generate Word...",
        command=lambda: self.generate("docx")
    ).pack(side="left", padx=(0, 7))
    ttk.Button(
        actions, text="Preview & Save PDF...",
        command=self.preview_pdf
    ).pack(side="left", padx=(0, 7))
    ttk.Button(
        actions, text="Generate Word + PDF",
        command=lambda: self.generate("both")
    ).pack(side="left", padx=(0, 7))
    ttk.Button(
        actions, text="Export Selected Data (ZIP/CSV)...",
        command=self.export_selected_professional_data
    ).pack(side="left", padx=(0, 7))
    ttk.Button(
        actions, text="Open Resume Folder",
        command=self.open_resume_folder
    ).pack(side="right")


def _options_from_app(self) -> dict[str, bool]:
    return {k: bool(v.get()) for k, v in self.cv_options.items()}


def enhanced_preview_pdf(self):
    app_module = __import__(self.__class__.__module__)
    profile = self.db.get_profile()
    base = (
        profile.get("preferred_name") or profile.get("full_name") or "Forensic"
    ).replace(" ", "_") + "_CV.pdf"
    temp_path = app_module.make_preview_temp_path()
    try:
        options = _options_from_app(self)
        self.set_status("Rendering selected PDF output...")
        self.update_idletasks()
        generate_extended_pdf(
            self.db, temp_path, options, theme_name="Professional"
        )
        default_save = app_module.portable_resume_dir() / base
        app_module.PdfPreviewWindow(
            self, temp_path, default_save,
            on_saved=lambda p: self.set_status(f"PDF saved: {p}")
        )
        self.set_status("PDF preview ready")
    except Exception as exc:
        messagebox.showerror("PDF Preview", str(exc), parent=self)
        self.set_status("Ready")


def enhanced_generate(self, output_type: str = "docx"):
    app_module = __import__(self.__class__.__module__)
    profile = self.db.get_profile()
    base = (
        profile.get("preferred_name") or profile.get("full_name") or "Forensic"
    ).replace(" ", "_") + "_CV"

    extension = ".pdf" if output_type == "pdf" else ".docx"
    out = filedialog.asksaveasfilename(
        title="Save Selected CV / Professional Output",
        defaultextension=extension,
        initialdir=str(app_module.portable_resume_dir()),
        initialfile=base + extension,
        filetypes=(
            [("PDF Document", "*.pdf")]
            if extension == ".pdf"
            else [("Word Document", "*.docx")]
        ),
    )
    if not out:
        return

    try:
        options = _options_from_app(self)
        if output_type == "pdf":
            generated = generate_extended_pdf(
                self.db, Path(out), options, theme_name="Professional"
            )
        else:
            docx_path = Path(out)
            generated = generate_extended_docx(self.db, docx_path, options)
            if output_type == "both":
                generate_extended_pdf(
                    self.db, docx_path.with_suffix(".pdf"), options,
                    theme_name="Professional",
                )
        self.set_status(f"Output generated: {generated}")
        if messagebox.askyesno(
            "Output Generated",
            "The selected output was created successfully. Open the output folder now?",
            parent=self,
        ):
            self.open_resume_folder()
    except Exception as exc:
        messagebox.showerror("Generation Error", str(exc), parent=self)


def export_selected_professional_data(self):
    app_module = __import__(self.__class__.__module__)
    profile = self.db.get_profile()
    base = (
        profile.get("preferred_name") or profile.get("full_name") or "Forensic"
    ).replace(" ", "_") + "_Selected_Data.zip"
    out = filedialog.asksaveasfilename(
        title="Export Selected Detailed Data",
        defaultextension=".zip",
        initialdir=str(app_module.portable_resume_dir()),
        initialfile=base,
        filetypes=[("ZIP Archive", "*.zip")],
    )
    if not out:
        return
    try:
        options = _options_from_app(self)
        path = export_selected_data_zip(self.db, out, options)
        self.set_status(f"Selected data exported: {path}")
        messagebox.showinfo(
            "Data Export Complete",
            "Selected detailed records were exported to CSV files inside the ZIP archive.",
            parent=self,
        )
    except Exception as exc:
        messagebox.showerror("Data Export", str(exc), parent=self)


def enhanced_refresh_dashboard(self):
    app_module = __import__(self.__class__.__module__)
    for w in self.metrics_frame.winfo_children():
        w.destroy()

    m = professional_metrics(self.db)
    metrics = [
        ("Cases", f"{m['casework']['unique_cases']:,}"),
        ("Examinations", f"{m['casework']['examinations']:,}"),
        ("Exam Hours", f"{m['casework']['hours']:,.1f}"),
        ("Reports", f"{m['reports']['total']:,}"),
        ("Peer Reviews", f"{m['reviews']['total']:,}"),
        ("Testimony", f"{m['court']['testimony_total']:,}"),
        ("Training Hrs", f"{m['training']['hours']:,.1f}"),
        ("Certifications", f"{m['training']['certifications']:,}"),
    ]
    for i, (label, value) in enumerate(metrics):
        box = ttk.LabelFrame(self.metrics_frame, text=label, padding=8)
        box.grid(row=i // 4, column=i % 4, padx=4, pady=4, sticky="nsew")
        ttk.Label(box, text=value, style="Metric.TLabel").pack()
        self.metrics_frame.columnconfigure(i % 4, weight=1)

    alerts = []
    try:
        expiring = self.db.conn.execute(
            "SELECT certification, expiration_date FROM certifications "
            "WHERE profile_id=? AND expiration_date <> '' ORDER BY expiration_date",
            (self.db.current_profile_id,),
        ).fetchall()
        for r in expiring:
            raw = str(r["expiration_date"] or "")
            label = "review date"
            try:
                normalized = normalize_date(raw, allow_present=False, year_only_ok=True)
                parts = normalized.split("-")
                y = int(parts[0])
                month = int(parts[1]) if len(parts) > 1 else 12
                day = int(parts[2]) if len(parts) > 2 else 28
                days = (date(y, month, min(day, 28)) - date.today()).days
                if days < 0:
                    label = "EXPIRED"
                elif days <= 365:
                    label = f"expires in {days} days"
                else:
                    label = "active"
            except Exception:
                pass
            alerts.append(f"• {r['certification']}: {raw} — {label}")
    except Exception:
        pass

    c = m["casework"]
    summary_lines = [
        "Professional record snapshot",
        "",
        f"Unique cases: {c['unique_cases']:,}",
        f"Forensic examinations: {c['examinations']:,}",
        f"Forensic reports tracked: {m['reports']['total']:,}",
        f"Peer / technical reviews: {m['reviews']['total']:,}",
        f"Court qualification records: {m['court']['qualification_records']:,}",
        f"Expert-witness testimony: {m['court']['expert']:,}",
        f"Training hours: {m['training']['hours']:,.1f}",
        f"CPE / CE credits: {m['training']['cpe']:,.1f}",
        f"Instruction hours: {m['teaching']['hours']:,.1f}",
        f"Presentations: {m['professional_output']['presentations']:,}",
        f"Publications / research: {m['professional_output']['publications']:,}",
        f"Tool validations: {m['quality']['validations']:,}",
        f"SOP / policy records: {m['quality']['procedures']:,}",
        f"Mentoring records: {m['leadership']['mentoring']:,}",
        f"Major projects: {m['leadership']['projects']:,}",
    ]
    if c["data_examined"]:
        summary_lines.append(f"Documented data examined: {format_bytes(c['data_examined'])}")

    summary_lines.extend([
        "",
        "Certification expiration alerts",
        "",
        *(alerts if alerts else ["No expiration dates entered."]),
        "",
        "Portable database",
        str(self.db_path),
        "",
        "Generated CV folder",
        str(app_module.portable_resume_dir()),
    ])

    self.summary_text.config(state="normal")
    self.summary_text.delete("1.0", "end")
    self.summary_text.insert("1.0", "\n".join(summary_lines))
    self.summary_text.config(state="disabled")

    counts = [
        ("Case Work", c["examinations"]),
        ("Reports", m["reports"]["total"]),
        ("Reviews", m["reviews"]["total"]),
        ("Testimony", m["court"]["testimony_total"]),
        ("Training", m["training"]["records"]),
        ("Certs", m["training"]["certifications"]),
        ("Teaching", m["teaching"]["records"]),
        ("Publications", m["professional_output"]["publications"]),
    ]
    self.dashboard_chart.delete("all")
    self.dashboard_chart.update_idletasks()
    width = max(self.dashboard_chart.winfo_width(), 360)
    height = max(self.dashboard_chart.winfo_height(), 260)
    max_value = max([v for _, v in counts] + [1])
    left, top, bottom = 78, 12, height - 22
    usable = max(120, width - left - 35)
    bar_h = max(12, min(26, int((bottom - top) / len(counts) - 5)))
    chart_fg = "#f2f3f5" if self.theme_name == "dark" else "#212529"
    for i, (label, value) in enumerate(counts):
        y = top + i * ((bottom - top) / len(counts))
        x2 = left + usable * (value / max_value)
        self.dashboard_chart.create_text(
            left - 7, y + bar_h / 2, text=label, anchor="e", fill=chart_fg
        )
        self.dashboard_chart.create_rectangle(
            left, y, x2, y + bar_h, fill="#1f4e79", outline=""
        )
        self.dashboard_chart.create_text(
            x2 + 5, y + bar_h / 2, text=str(value), anchor="w", fill=chart_fg
        )


_APP_INSTALLED = False


def install_app_extensions(App) -> None:
    """Install enhanced UI/output methods after App class has been defined."""
    global _APP_INSTALLED
    if _APP_INSTALLED:
        return

    app_module = __import__(App.__module__)
    # RecordsTab resolves RecordDialog from app module globals at runtime.
    setattr(app_module, "RecordDialog", ScrollableRecordDialog)

    App._generate_tab = enhanced_generate_tab
    App.preview_pdf = enhanced_preview_pdf
    App.generate = enhanced_generate
    App.export_selected_professional_data = export_selected_professional_data
    App.refresh_dashboard = enhanced_refresh_dashboard
    _APP_INSTALLED = True
